from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()

# Add Title
title = doc.add_heading('Struktur Fitur Sistem Logistik Rumah Sakit', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph('Berikut adalah daftar fitur dan modul komprehensif untuk Sistem Logistik Rumah Sakit, dikelompokkan berdasarkan Role (Peran) masing-masing pengguna.')

# Role 1
doc.add_heading('1. ROLE: ADMIN SISTEM', level=1)
doc.add_paragraph('Peran ini berfokus pada pengaturan dasar aplikasi, manajemen master data, keamanan, dan workflow rumah sakit.')

doc.add_heading('Modul Manajemen Pengguna & Hak Akses (User Management)', level=2)
doc.add_paragraph('• Kelola akun: Tambah, edit, hapus, dan reset password untuk petugas Logistik dan staf Unit.\n'
                  '• Role & Permission: Mengatur batasan akses (misal: Siapa Kepala Logistik yang bisa approve pengadaan, siapa yang hanya bisa melihat).', style='List Bullet')

doc.add_heading('Modul Master Data (Core Database)', level=2)
doc.add_paragraph('• Master Data Unit: Mengelola daftar unit kerja/ruangan (VK, IGD, CSSD, Rawat Inap).\n'
                  '• Master Data Kategori: Mengatur kategori barang (Alkes, APD, ATK, Cetakan, Alat Medis, dll).\n'
                  '• Master Data Barang (Katalog): Database detail seluruh barang yang mencakup nama, satuan (Pcs, Box, Rim), harga dasar, dan Batas Stok Minimum (Low Stock Limit).\n'
                  '• Master Data Supplier (Rekanan): Daftar vendor penyedia/distributor barang logistik.\n'
                  '• Master Data Lokasi (Rak/Gudang): Penamaan lokasi rak di gudang agar petugas cepat mencari barang.', style='List Bullet')

doc.add_heading('Modul Konfigurasi & Keamanan', level=2)
doc.add_paragraph('• Log Aktivitas (Audit Trail): Memantau semua histori aktivitas demi keamanan (Anti-Fraud).\n'
                  '• Sistem Notifikasi: Konfigurasi pop-up/email/WA untuk memberitahu Logistik jika ada permintaan Urgent/Cito.', style='List Bullet')


# Role 2
doc.add_heading('2. ROLE: PETUGAS LOGISTIK UMUM', level=1)
doc.add_paragraph('Peran ini adalah eksekutor yang memproses permintaan unit, mengelola gudang, dan melakukan pengadaan ke pihak luar.')

doc.add_heading('A. Modul Manajemen Permintaan Masuk (Order Fulfillment)', level=2)
doc.add_paragraph('• Daftar Antrean Permintaan: Menampilkan tiket permintaan (Rutin, Urgent/Cito, Pengadaan Baru).\n'
                  '• Proses & Validasi Item (Fitur Checklist):\n'
                  '  - Tandai Siap: Barang tersedia penuh.\n'
                  '  - Tandai Kosong: Stok gudang habis sama sekali.\n'
                  '  - Ubah Jumlah (Parsial): Minta 10, disetujui 5 karena stok menipis.\n'
                  '• Serah Terima Barang: Tombol/Tanda tangan digital konfirmasi bahwa perawat unit sudah mengambil barang dari gudang.\n'
                  '• Retur Barang Internal: Fitur menerima pengembalian barang dari Unit (karena salah spesifikasi, cacat, atau kedaluwarsa).', style='List Bullet')

doc.add_heading('B. Modul Manajemen Stok & Inventaris Gudang (Inventory Control)', level=2)
doc.add_paragraph('• Penerimaan Barang (Goods Receipt): Form input saat barang dari Supplier datang. Otomatis menambah stok di gudang.\n'
                  '• Kartu Stok (Stock Card): Pencatatan riwayat real-time pergerakan tiap item (Stok Awal -> Masuk -> Keluar -> Saldo Akhir).\n'
                  '• Manajemen Kadaluarsa (FEFO/FIFO): Peringatan untuk barang medis/cairan yang mendekati Expired Date.\n'
                  '• Stok Opname: Fitur penyesuaian/sinkronisasi antara stok fisik di lapangan vs stok di sistem.\n'
                  '• Inventaris Aset & Maintenance Tracking: Data aset RS beserta histori kondisinya, dan status perbaikan vendor.', style='List Bullet')

doc.add_heading('C. Modul Manajemen Pengadaan (Purchasing)', level=2)
doc.add_paragraph('• Daftar Ajuan Pengadaan: Pembuatan draft usulan belanja ke manajemen/keuangan jika stok menipis atau kosong.\n'
                  '• Pembuatan Purchase Order (PO): Cetak surat pesanan resmi dari RS untuk dikirimkan ke Supplier.\n'
                  '• Penerimaan Faktur (Invoice): Pencatatan tagihan dari supplier saat barang datang untuk diteruskan ke Keuangan.', style='List Bullet')

doc.add_heading('D. Modul Laporan & Dasbor Analitik', level=2)
doc.add_paragraph('• Dasbor Visual (Chart): Grafik tren permintaan, unit paling boros, dan Service Level.\n'
                  '• Rekapitulasi: Laporan distribusi per unit per minggu/bulan.\n'
                  '• Laporan Stok Kritis & Dead Stock: Laporan barang yang sudah mau habis, dan barang yang tidak pernah diminta unit.', style='List Bullet')

# Role 3
doc.add_heading('3. ROLE: UNIT RS (IGD, VK, CSSD, Rawat Inap, dll.)', level=1)
doc.add_paragraph('Peran ini bertindak sebagai pemohon (Requester) barang dan perbaikan.')

doc.add_heading('A. Modul Pengajuan Permintaan (Requisition)', level=2)
doc.add_paragraph('• Permintaan Rutin Mingguan: Keranjang belanja barang operasional harian. Dijadwalkan otomatis pengambilannya.\n'
                  '• Permintaan Urgent (Cito): Form khusus untuk barang kritis. Langsung berbunyi/berkedip di layar Logistik.\n'
                  '• Permintaan Pengadaan (Barang Baru): Request alat atau barang yang belum ada di katalog (Master Data).', style='List Bullet')

doc.add_heading('B. Modul Pelaporan & Monitoring (Tracking)', level=2)
doc.add_paragraph('• Histori Permintaan (Live Tracking): Lacak tiket permintaan (Menunggu -> Diproses -> Siap Diambil -> Selesai).\n'
                  '• Edit/Batal (Draft): Bisa ubah pesanan selama status masih "Menunggu".\n'
                  '• Laporan Barang Rusak: Pengajuan maintenance atau servis ke teknisi logistik.\n'
                  '• Konfirmasi Terima Barang: Tombol otorisasi di pihak Unit yang menyatakan barang sudah diterima fisik.\n'
                  '• Stok Depo Ruangan (Opsional): Menampilkan sisa stok barang di lemari ruangan mereka sendiri.', style='List Bullet')


doc.add_heading('Alur Kerja (Workflow) yang Disempurnakan:', level=1)
doc.add_paragraph('1. Siklus Normal (Distribusi Internal): Unit ajukan Permintaan Rutin -> Logistik siapkan barang -> Unit datang ambil barang -> Unit klik "Konfirmasi Terima" -> Kartu Stok gudang otomatis berkurang.\n'
                  '2. Siklus Restock (Pengadaan Eksternal): Sistem mendeteksi stok Kritis -> Muncul peringatan di Logistik -> Logistik buat PO ke Supplier -> Barang datang dari Supplier -> Logistik input "Penerimaan Barang" -> Kartu Stok gudang otomatis bertambah.\n'
                  '3. Siklus Kekosongan: Unit minta barang -> Logistik tandai "Kosong" -> Item otomatis masuk Draft Pengadaan -> Unit bisa memantau pesanan.\n'
                  '4. Siklus Aset/Servis: Unit lapor aset rusak -> Logistik mengubah status inventaris ke "Sedang Diservis Vendor" -> Setelah selesai, status dikembalikan ke "Baik" -> Unit mendapat notifikasi.', style='List Number')

doc.save('Rancangan_Sistem_Logistik_RS.docx')
print("Document saved as Rancangan_Sistem_Logistik_RS.docx")
