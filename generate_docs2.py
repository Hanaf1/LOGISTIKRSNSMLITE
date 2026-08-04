from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def add_table(doc, headers, rows, header_bg=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0] if hdr_cells[i].paragraphs[0].runs else hdr_cells[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        if header_bg:
            tc = hdr_cells[i]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), header_bg)
            tcPr.append(shd)
    for ri, row_data in enumerate(rows):
        row_cells = table.rows[ri+1].cells
        for ci, cell_val in enumerate(row_data):
            row_cells[ci].text = str(cell_val)
            for para in row_cells[ci].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    return table

def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    p.paragraph_format.left_indent = Inches(0.3)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    return p

def note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run('\u26a0 ' + text)
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = RGBColor(0xB8, 0x56, 0x00)
    return p

# ========================
doc = Document()
section = doc.sections[0]
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

# === JUDUL ===
title = doc.add_heading('DOKUMENTASI PERUBAHAN SISTEM', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x15, 0x63, 0xBE)

sub = doc.add_paragraph('Modul Logistik Non Medis \u2014 RSNS mLITE')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(11)
sub.runs[0].bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Kode Sumber Awal (Baseline): ').bold = True
p.add_run('https://github.com/Rifangga99/mlite_rsns.git')
p.add_run(' (branch main, commit terakhir: 027fcfd)')

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.add_run('Kode Setelah Penyesuaian: ').bold = True
p2.add_run('https://github.com/Hanaf1/LOGISTIKRSNSMLITE.git')

doc.add_paragraph(f'Tanggal Dokumentasi: {datetime.datetime.now().strftime("%d %B %Y")}').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

note(doc, 'Dokumen ini menjelaskan semua perubahan yang dilakukan SETELAH mengambil kode dari repo Rifangga. '
     'Perubahan mencakup: penyesuaian menu & fitur, perubahan struktur tabel database, dan konfigurasi notifikasi Workerman.')

doc.add_paragraph()
doc.add_heading('DAFTAR ISI', level=1)
daftar_isi = [
    ('1', 'PENYESUAIAN MENU & FITUR ANTARMUKA'),
    ('  1.1', 'Menu / Halaman Baru yang Ditambahkan'),
    ('  1.2', 'Menu / Fitur yang Diubah dari Rifangga'),
    ('  1.3', 'Fitur yang Dihapus'),
    ('  1.4', 'Plugin Baru: permintaan_logistik_non_medis'),
    ('2', 'PENYESUAIAN TABEL DATABASE'),
    ('  2.1', 'Tabel sppb \u2014 Kolom Baru & Perubahan ENUM'),
    ('  2.2', 'Tabel aset \u2014 Kolom Baru untuk KIB'),
    ('  2.3', 'Tabel master_barang \u2014 Kolom Baru'),
    ('  2.4', 'Tabel perencanaan \u2014 Kolom Baru'),
    ('  2.5', 'Tabel pr (Purchase Request) \u2014 Perubahan ENUM'),
    ('  2.6', 'Tabel inventaris_master \u2014 Kolom & Index Baru'),
    ('  2.7', 'Tabel Baru yang Dibuat'),
    ('3', 'PENGATURAN NOTIFIKASI WORKERMAN'),
    ('  3.1', 'Perubahan workerman.php dari Baseline Rifangga'),
    ('  3.2', 'Cara Kerja Queue File'),
    ('  3.3', 'Cara Menjalankan'),
]
for no, judul in daftar_isi:
    p = doc.add_paragraph(f'{no}   {judul}')
    p.paragraph_format.space_after = Pt(2)
    if len(no.strip()) == 1:
        p.runs[0].bold = True

doc.add_page_break()

# ==============================================================================
# BAGIAN 1: PENYESUAIAN MENU
# ==============================================================================
doc.add_heading('1. PENYESUAIAN MENU & FITUR ANTARMUKA', level=1)
doc.add_paragraph(
    'Perubahan ini dilakukan di atas kode Rifangga. Beberapa halaman baru ditambahkan, '
    'beberapa fitur diubah, dan satu fitur dihapus.'
)

# 1.1 Menu Baru
doc.add_heading('1.1 Menu / Halaman Baru yang Ditambahkan', level=2)
note(doc, 'File-file berikut TIDAK ADA di repo Rifangga, dibuat baru oleh tim pengembang RSNS.')
doc.add_paragraph()

add_table(doc, ['Nama Menu', 'URL / Endpoint', 'File View (Baru)', 'Fungsi'],
[
    ['Master Inventaris Barang', '/logistik_non_medis/masterinventaris', 'master.inventaris.html\nmaster.inventaris.barang.display.html', 'Kelola hierarki Kategori-Kelompok-Jenis master barang inventaris non-medis'],
    ['Kategori Aset KIB', '/logistik_non_medis/masterkategoriaset', 'master_kategori_aset.html', 'Kelola kategori KIB: A=Tanah, B=Mesin/Peralatan, C=Kendaraan, D=Gedung, E=Aset Tetap Lain, F=Konstruksi'],
    ['Notifikasi Real-Time', '/logistik_non_medis/notifikasi', 'notifikasi.html', 'Halaman daftar notifikasi masuk per user (SPPB baru, Siap Ambil, dll)'],
    ['Opname Gudang V2', '/logistik_non_medis/gudangopnamev2', 'gudang.opnamev2.html\ngudang.opnamev2.display.html', 'Versi baru opname stok yang lebih ringkas dan mudah digunakan'],
    ['Distribusi BHP', '/logistik_non_medis/distribusibhp', 'distribusi.bhp.html\ndistribusi.bhp.display.html\ndistribusi.bhp.detail.html', 'Halaman distribusi Bahan Habis Pakai (BHP) tersendiri'],
    ['Cetak SPPB', '/logistik_non_medis/distribusisppbcetak', 'distribusi.sppb.cetak.html', 'Halaman cetak/print dokumen SPPB'],
    ['Realisasi Belanja', '/logistik_non_medis/realisasibelanja', 'pengadaan.realisasi.html', 'Monitoring dan pelaporan realisasi anggaran belanja pengadaan'],
    ['Permintaan Non Rutin (Form Unit)', '/logistik_non_medis/distribusinonrutinform', 'pengadaan.permintaan_nonrutin.html', 'Form khusus untuk unit mengajukan permintaan barang non-rutin'],
    ['Rencana Rutin', '/logistik_non_medis/rencanabelanja / rencanarutin', 'pengadaan.rencana_rutin.html\npengadaan.rencana_rutin.form.html\npengadaan.rencana_rutin.detail.html', 'Perencanaan kebutuhan barang rutin per periode'],
    ['Rencana Non Rutin', '/logistik_non_medis/rencanabelanjanonrutin', 'pengadaan.rencana_nonrutin.html\npengadaan.rencana_nonrutin.form.html\npengadaan.rencana_nonrutin.detail.html', 'Perencanaan kebutuhan barang non-rutin'],
    ['Rencana Pembelian', '/logistik_non_medis/rencanapembelian', 'pengadaan.rencana_pembelian.html', 'Daftar rencana pembelian berdasarkan kebutuhan yang disetujui'],
    ['Terima Rutin', '/logistik_non_medis/terimabarangrutin', 'pengadaan.terima_rutin.html\npengadaan.terima_rutin.form.html\npengadaan.terima_rutin.detail.html', 'Penerimaan dan verifikasi barang rutin dari supplier/gudang pusat'],
    ['Laporan Aset', '/logistik_non_medis/laporanaset', 'laporan.aset.html', 'Laporan rekap dan export data aset KIB'],
    ['Laporan Cost Unit', '/logistik_non_medis/laporancostunit', 'laporan.costunit.html\nlaporan.costunit.display.html\nlaporan.costunit.detail.html', 'Laporan biaya/cost per unit pelayanan rumah sakit'],
    ['Laporan Pengadaan', '/logistik_non_medis/laporanpengadaan', 'laporan.pengadaan.html', 'Laporan rekapitulasi pengadaan barang'],
    ['Pemeliharaan Aset (Dashboard)', '/logistik_non_medis/asetpemeliharaan', 'aset.pemeliharaan.dashboard.html', 'Dashboard jadwal dan status pemeliharaan aset'],
    ['Detail Registrasi Aset', '/logistik_non_medis/asetregistrasidetail', 'aset.registrasi.detail.html', 'Halaman detail view satu aset (read-only)'],
    ['Detail Mutasi Aset', '/logistik_non_medis/asetmutasidetail', 'aset.mutasi.detail.html\naset.mutasi.form.html', 'Detail dan form mutasi/perpindahan aset antar unit'],
], header_bg='D6F7E4')

# 1.2 Menu yang Diubah
doc.add_heading('1.2 Menu / Fitur yang Diubah dari Baseline Rifangga', level=2)
note(doc, 'File-file berikut ADA di repo Rifangga tetapi dimodifikasi oleh tim pengembang RSNS.')
doc.add_paragraph()

add_table(doc, ['Halaman / Fitur', 'File yang Diubah', 'Perubahan yang Dilakukan'],
[
    ['Registrasi Aset (List)', 'aset.registrasi.html\naset.registrasi.display.html', 'Perbaikan pagination (dari bug div di tbody menjadi JSON response). Data yang tampil sebelumnya cuma 5, sekarang tampil lengkap dengan paginasi berjalan.'],
    ['Form Registrasi Aset', 'aset.registrasi.form.html', 'Penambahan dropdown Kategori \u2192 Kelompok \u2192 Jenis yang terhubung ke tabel master inventaris. Sebelumnya tidak ada pemilihan kategori.'],
    ['Aset KIB (List & Form)', 'aset.kib.html\naset.kib.display.html\naset.kib.rekap.html', 'Penambahan kolom kib_jenis (golongan A-F), kib_nilai_buku, kib_akumulasi_penyusutan, kib_tgl_mulai, nomor_inventaris, merk_type, lokasi_fisik ke tampilan dan form.'],
    ['Distribusi SPPB (Rutin)', 'distribusi.sppb.html\ndistribusi.sppb.display.html\ndistribusi.sppb.form.html\ndistribusi.sppb.detail.html', '1) Tab "Aktif" & "Riwayat" diperbaiki sehingga memfilter status untuk SEMUA role (Admin & Logistik), bukan hanya unit.\n2) Perluasan alur persetujuan: tambahan status "Disetujui Ka. Sie", "Disetujui Kabid", "Siap Diserahkan", dll.\n3) Form SPPB mendukung pemilihan jenis Rutin vs Non Rutin.'],
    ['Dashboard Manage', 'manage.html', 'Penambahan widget counter baru (SPPB minggu ini, Pending Verifikasi, Proses Packing). Tab menu disesuaikan dengan role user. Badge notifikasi SPPB pending muncul.'],
    ['Form Unit', 'master.unit.html\nmaster.unit.form.html\nmaster.unit.display.html', 'Saat PJ (Penanggung Jawab) unit diubah, data user_roles otomatis disinkronkan. Sebelumnya ada bug di sini.'],
    ['Hak Akses', 'hakakses.html\nhakakses.form.html', 'Penyesuaian tampilan dan perbaikan form mapping user ke role dan unit.'],
    ['Perencanaan', 'pengadaan.perencanaan.html\nperencanaan.form.html', 'Penambahan field kelompok_barang pada form. Penyesuaian tampilan untuk mendukung rencana rutin dan non-rutin.'],
    ['PO (Purchase Order)', 'pengadaan.po.html\npengadaan.po.form.html\npengadaan.po.display.html\npengadaan.po.cetak.html', 'Penambahan kolom estimasi_harga, update form PO dan tampilan cetakan.'],
    ['Gudang Stok', 'gudang.stok.html\ngudang.stok.display.html', 'Penyesuaian tampilan, filter, dan export stok gudang.'],
    ['Gudang Mutasi', 'gudang.mutasi.html\ngudang.mutasi.form.html\ngudang.mutasi.display.html', 'Perbaikan form mutasi dan tampilan riwayat mutasi barang.'],
    ['Master Barang', 'master.barang.html\nmaster.barang.form.html\nmaster.barang.display.html\nmaster.barang.detail.html', 'Penambahan field kategori dan kode_kategori pada master barang. Perbaikan handler penghapusan.'],
    ['Laporan Distribusi', 'laporan.distribusi.html', 'Perluasan filter dan tampilan laporan distribusi SPPB.'],
    ['Serah Terima (BAST)', 'distribusi.serahterima.form.html\ndistribusi.serahterima.bast.html', 'Perubahan format dokumen Berita Acara Serah Terima (BAST).'],
    ['Laporan Stok Mutasi', 'laporan.stokmutasi.html\nlaporan.eksporcetak.html', 'Penyesuaian tampilan dan tambahan opsi filter.'],
    ['Workerman (WebSocket)', 'workerman.php', 'Penambahan Timer Queue-based broadcasting. Sebelumnya hanya echo balik dari browser. Lihat Bagian 3 untuk detail lengkap.'],
], header_bg='FFF3CD')

# 1.3 Fitur Dihapus
doc.add_heading('1.3 Fitur yang Dihapus', level=2)

add_table(doc, ['Fitur', 'Lokasi Sebelumnya', 'Alasan Dihapus'],
[
    ['Tombol "Import SPPB Mingguan"', 'distribusi.sppb.html (baris 84)', 'Dihapus berdasarkan permintaan pengguna. Fitur import mingguan tidak dipakai.'],
    ['Modal Import SPPB Mingguan', 'distribusi.sppb.html (modal baris 283)', 'Dihapus bersama tombol di atas.'],
    ['Kolom minggu_ke di form SPPB', 'distribusi.sppb.form.html', 'Tidak relevan setelah fitur import mingguan dihapus.'],
], header_bg='FFD6D6')

# 1.4 Plugin Baru
doc.add_heading('1.4 Plugin Baru: permintaan_logistik_non_medis', level=2)
doc.add_paragraph(
    'Plugin baru ini TIDAK ADA di repo Rifangga. Dibuat sebagai modul terpisah untuk memungkinkan '
    'user/unit mengajukan permintaan barang tanpa perlu akses ke modul logistik_non_medis penuh.'
)

add_table(doc, ['File', 'Fungsi'],
[
    ['plugins/permintaan_logistik_non_medis/Admin.php', 'Controller utama modul permintaan. Berisi endpoint untuk list, form, tambah, dan detail permintaan.'],
    ['plugins/permintaan_logistik_non_medis/Info.php', 'Informasi plugin (nama, versi, deskripsi) untuk sistem mLITE.'],
    ['view/admin/manage.html', 'Halaman utama list permintaan untuk user/unit.'],
    ['view/admin/form.html', 'Form pengajuan permintaan barang baru.'],
    ['view/admin/detail.html', 'Halaman detail permintaan yang sudah diajukan.'],
    ['view/admin/settings.html', 'Halaman pengaturan plugin.'],
], header_bg='D6E4F7')

doc.add_page_break()

# ==============================================================================
# BAGIAN 2: PENYESUAIAN TABEL DATABASE
# ==============================================================================
doc.add_heading('2. PENYESUAIAN TABEL DATABASE', level=1)
doc.add_paragraph(
    'Semua tabel menggunakan awalan rsns_custom_logistik_non_medis_. '
    'Perubahan di bawah ini adalah perubahan DARI kondisi tabel di repo Rifangga.'
)

# 2.1 sppb
doc.add_heading('2.1 Tabel: sppb', level=2)
note(doc, 'Di repo Rifangga: tabel sppb sudah ada, namun beberapa kolom dan nilai ENUM berbeda.')

add_table(doc, ['Jenis', 'Kolom', 'Detail Perubahan'],
[
    ['ALTER MODIFY', 'status (ENUM)', 'Ditambahkan status baru:\n- "Diserahkan ke Kasie Umum"\n- "Verifikasi Kasie Umum"\n- "Diteruskan ke Logistik Umum"\n- "Rekap Logistik"\n- "Logistik Umum & Rekap"\n- "Konsultasi Dana"\n- "Konsul Pengajuan ke Kabid Umum"\n- "Diserahkan ke Keuangan"\n- "Pengajuan Dana ke Bendahara"\n- "Tidak ACC"\n- "Siap Diserahkan"\n- "Dibatalkan"\n\nPada versi Rifangga tidak ada status-status ini.'],
    ['ALTER ADD', 'ditolak_pada_status', 'varchar(100) — menyimpan status SPPB saat ditolak, agar bisa diaudit tanpa JOIN tabel log.\nPosisi: AFTER alasan_penolakan'],
    ['ALTER ADD', 'jenis_permintaan', "varchar(50) DEFAULT 'Rutin' — membedakan alur Rutin vs Non Rutin.\nPosisi: AFTER kode_unit\nData lama: diisi 'Non Rutin' jika user_approve_ka_unit tidak kosong."],
    ['ALTER ADD', 'minggu_ke', 'int(11) DEFAULT NULL — nomor minggu dalam tahun untuk import mingguan.\nPosisi: AFTER tgl_sppb'],
    ['ALTER ADD', 'estimasi_harga', 'double NOT NULL DEFAULT 0 — estimasi total harga SPPB.\nPosisi: AFTER spesifikasi_manual'],
    ['Kolom existing (Rifangga)', 'user_approve_ka_unit\nuser_approve_ka_sie\nuser_approve_ka_bidang\ndiambil_oleh, tgl_diambil', 'Kolom-kolom ini sudah ada di Rifangga dan dipertahankan. Tidak diubah.'],
], header_bg='D6E4F7')

# 2.2 aset
doc.add_heading('2.2 Tabel: aset (Registrasi & KIB)', level=2)
note(doc, 'Di repo Rifangga: tabel aset ada tetapi tidak memiliki kolom-kolom KIB dan inventaris di bawah ini.')

add_table(doc, ['Jenis', 'Kolom Baru', 'Tipe Data', 'Keterangan'],
[
    ['ALTER ADD', 'nomor_inventaris', 'varchar(50)', 'Nomor inventaris manual (selain kode_aset otomatis). AFTER kode_aset'],
    ['ALTER ADD', 'merk_type', 'varchar(150)', 'Merk / tipe aset (contoh: "HP LaserJet Pro M428"). AFTER nama_aset'],
    ['ALTER ADD', 'nomor_dokumen', 'varchar(200)', 'Nomor dokumen pembelian/berita acara. AFTER serial_number'],
    ['ALTER ADD', 'lokasi_fisik', 'varchar(150)', 'Lokasi fisik aset dipasang/disimpan. AFTER kode_lokasi'],
    ['ALTER ADD', 'bahan', 'varchar(100)', 'Bahan/material utama aset. AFTER lokasi_fisik'],
    ['ALTER ADD', 'tahun_beli', 'smallint(4)', 'Tahun perolehan aset. AFTER tanggal_perolehan'],
    ['ALTER ADD', 'satuan', 'varchar(50)', "Satuan aset ('Buah','Unit','Set'). AFTER tahun_beli"],
    ['ALTER ADD', 'jumlah', 'int(11) DEFAULT 1', 'Jumlah fisik aset. AFTER satuan'],
    ['ALTER ADD', 'keterangan_inventaris', 'text', 'Catatan kondisi/keterangan khusus. AFTER status_kondisi'],
    ['ALTER ADD KEY', 'INDEX nomor_inventaris', '-', 'Index baru untuk pencarian cepat berdasarkan nomor inventaris'],
    ['Kolom KIB (baru)', 'kib_jenis', "ENUM('A','B','C','D','E','F')", 'Golongan KIB: A=Tanah, B=Peralatan Mesin, C=Gedung & Bangunan, D=Jalan/Irigasi/Jaringan, E=Aset Tetap Lainnya, F=Konstruksi dlm Pengerjaan'],
    ['Kolom KIB (baru)', 'kib_tgl_mulai', 'date', 'Tanggal mulai penggunaan aset dalam KIB'],
    ['Kolom KIB (baru)', 'kib_tgl_rencana_selesai', 'date', 'Tanggal rencana selesai/hapus dari KIB'],
    ['Kolom KIB (baru)', 'kib_nilai_buku', 'decimal(15,2)', 'Nilai buku aset saat ini setelah penyusutan'],
    ['Kolom KIB (baru)', 'kib_akumulasi_penyusutan', 'decimal(15,2)', 'Total akumulasi penyusutan yang sudah dicatat'],
], header_bg='D6F7E4')

# 2.3 master_barang
doc.add_heading('2.3 Tabel: master_barang', level=2)
note(doc, 'Di repo Rifangga: kolom kategori dan kode_kategori belum ada di tabel master_barang.')

add_table(doc, ['Jenis', 'Kolom Baru', 'Tipe Data', 'Keterangan'],
[
    ['ALTER ADD', 'kategori', 'varchar(100) DEFAULT NULL', 'Nama kategori barang (teks). AFTER spesifikasi'],
    ['ALTER ADD', 'kode_kategori', 'varchar(50) DEFAULT NULL', 'Kode kategori barang (relasi ke tabel inventaris_kategori). AFTER kategori'],
], header_bg='F7EAD6')

# 2.4 perencanaan
doc.add_heading('2.4 Tabel: perencanaan', level=2)
note(doc, 'Di repo Rifangga: kolom kelompok_barang belum ada.')

add_table(doc, ['Jenis', 'Kolom Baru', 'Tipe Data', 'Keterangan'],
[
    ['ALTER ADD', 'kelompok_barang', 'varchar(100) DEFAULT NULL', 'Kelompok barang untuk perencanaan. AFTER kode_unit'],
], header_bg='F7EAD6')

# 2.5 pr
doc.add_heading('2.5 Tabel: pr (Purchase Request)', level=2)
note(doc, 'Di repo Rifangga: ENUM status PR berbeda.')
add_table(doc, ['Jenis', 'Kolom', 'Detail'],
[
    ['ALTER MODIFY', 'status (ENUM)', "Sebelum (Rifangga): tidak ada nilai 'Selesai'\nSesudah: ENUM('Draft','Diajukan','Disetujui','Di-PO-kan','Ditolak','Selesai')"],
], header_bg='EAD6F7')

# 2.6 inventaris_master
doc.add_heading('2.6 Tabel: inventaris_master', level=2)
note(doc, 'Di repo Rifangga: tabel inventaris_master sudah ada tetapi belum punya kolom hierarki kategori.')

add_table(doc, ['Jenis', 'Kolom', 'Tipe Data', 'Keterangan'],
[
    ['ALTER ADD', 'kode_inventaris', 'char(3) DEFAULT NULL', 'Kode inventaris singkat (untuk entri UNIT). AFTER kode'],
    ['ALTER MODIFY', 'kode', 'varchar(50) NOT NULL', 'Diperpanjang agar bisa menampung kode hierarki panjang'],
    ['ALTER ADD', 'kode_kategori', 'char(1) NOT NULL DEFAULT ""', 'Kategori (1 karakter: kode inventaris_kategori). AFTER kode_inventaris'],
    ['ALTER ADD', 'nama_kelompok', 'varchar(150) DEFAULT NULL', 'Cache nama kelompok (denormalisasi untuk performa). AFTER kode_barang'],
    ['ALTER ADD', 'nama_jenis', 'varchar(150) DEFAULT NULL', 'Cache nama jenis (denormalisasi untuk performa). AFTER nama_kelompok'],
    ['ALTER DROP INDEX + ADD UNIQUE', 'jenis_kode (INDEX)', '-', 'Index unik diubah dari (jenis_master, kode) menjadi (jenis_master, kode_kategori, kode)'],
], header_bg='EAD6F7')

# 2.7 Tabel Baru
doc.add_heading('2.7 Tabel Baru yang Dibuat', level=2)
note(doc, 'Tabel-tabel di bawah ini TIDAK ADA di repo Rifangga. Semuanya dibuat oleh pengembang RSNS.')

add_table(doc, ['Nama Tabel', 'Fungsi', 'Kolom Utama'],
[
    ['notifikasi', 'Menyimpan notifikasi real-time per user. Dibaca oleh halaman notifikasi.html dan di-push via Workerman.', 'id (PK), user_target (varchar), tipe (varchar), pesan (text), url (varchar), is_read (tinyint), tgl_dibuat (datetime)'],
    ['inventaris_kategori', 'Hierarki level 1: Kategori master barang inventaris (contoh: "1"=Inventaris, "2"=BHP).', 'kode_kategori (char 1, PK), nama_kategori, kib_default (ENUM A-F), umur_manfaat_default (int tahun), kode_coa (varchar), status_aktif'],
    ['inventaris_kelompok', 'Hierarki level 2: Kelompok dalam kategori (contoh: "01"=Peralatan Kesehatan).', 'id (PK), kode_kategori (FK), kode_kelompok (char 2), nama_kelompok'],
    ['inventaris_jenis', 'Hierarki level 3: Jenis dalam kelompok (contoh: "001"=Tensimeter Digital).', 'id (PK), kode_kategori (FK), kode_kelompok (FK), kode_jenis (char 3), nama_jenis'],
    ['penerimaan', 'Mencatat penerimaan barang dari supplier berdasarkan PO.', 'id (PK), no_penerimaan, tgl_penerimaan, no_po, kode_gudang, status, dibuat_oleh'],
    ['rencana_rutin', 'Rencana kebutuhan barang rutin per periode (bulanan/tahunan).', 'id (PK), no_rencana, tgl_rencana, kode_unit, periode, status'],
    ['rencana_rutin_detail', 'Detail item barang dalam rencana rutin.', 'id (PK), id_rencana (FK), kode_barang, nama_barang, qty_rencana, satuan'],
    ['rencana_nonrutin', 'Rencana kebutuhan barang non-rutin (pengadaan khusus).', 'id (PK), no_rencana, tgl_rencana, kode_unit, kebutuhan, status'],
    ['rencana_nonrutin_detail', 'Detail item barang dalam rencana non-rutin.', 'id (PK), id_rencana (FK), nama_barang, spesifikasi, qty, estimasi_harga'],
    ['terima_rutin', 'Penerimaan barang rutin dari gudang/distribusi.', 'id (PK), no_terima, tgl_terima, kode_unit, status, dibuat_oleh'],
    ['terima_rutin_detail', 'Detail item barang yang diterima pada penerimaan rutin.', 'id (PK), id_terima (FK), kode_barang, nama_barang, qty_terima, satuan'],
    ['opname_v2', 'Opname stok gudang versi 2 (lebih sederhana dari opname v1).', 'id (PK), no_opname, tgl_opname, kode_gudang, status, dibuat_oleh'],
], header_bg='EAD6F7')

doc.add_page_break()

# ==============================================================================
# BAGIAN 3: PENGATURAN NOTIFIKASI WORKERMAN
# ==============================================================================
doc.add_heading('3. PENGATURAN NOTIFIKASI WORKERMAN', level=1)
doc.add_paragraph(
    'Bagian ini menjelaskan perubahan sistem notifikasi real-time. '
    'Kode baseline dari Rifangga sudah menggunakan Workerman WebSocket, '
    'namun hanya bisa echo-balik pesan dari browser. Perubahan yang dilakukan '
    'adalah menambahkan Queue File System agar PHP backend bisa mengirim notifikasi ke browser.'
)

doc.add_heading('3.1 Perubahan workerman.php dari Baseline Rifangga', level=2)

add_table(doc, ['Bagian', 'Di Repo Rifangga (Sebelum)', 'Setelah Perubahan'],
[
    ['use statement', 'Hanya: use Workerman\\Worker;', 'Ditambah: use Workerman\\Timer;'],
    ['Queue File', 'Tidak ada', 'Variabel $queueFile = BASE_DIR/tmp/logistik_notifications.queue\nVariabel $queueOffset = 0 (tracking posisi baca)'],
    ['Broadcast Function', 'Tidak ada fungsi $broadcast terpisah', 'Fungsi $broadcast ditambahkan: kirim JSON ke semua koneksi WebSocket aktif'],
    ['onWorkerStart', 'Tidak ada', 'Ditambahkan event handler onWorkerStart yang menjalankan Timer setiap 0.5 detik untuk membaca baris baru dari queue file'],
    ['Timer Logic', 'Tidak ada', 'Timer: buka file dengan flock(LOCK_SH), fseek ke $queueOffset, baca baris baru, decode JSON, broadcast jika type=logistik_notification, update $queueOffset'],
    ['onMessage', 'Broadcast langsung ke semua koneksi', 'Tetap dipertahankan untuk backward-compatibility (pesan dari browser masih di-broadcast)'],
    ['onClose', 'Ada (log disconnect)', 'Tidak diubah'],
], header_bg='D6E4F7')

doc.add_heading('3.2 Cara Kerja Queue File (Alur Lengkap)', level=2)
add_code(doc, """ALUR PENGIRIMAN NOTIFIKASI:

[1] User/Admin melakukan aksi (misal: submit SPPB baru, atau ubah status menjadi "Siap Ambil")

[2] Controller PHP (Admin.php) menyimpan notifikasi ke tabel DB:
    $this->db('rsns_custom_logistik_non_medis_notifikasi')->save([
        'user_target' => 'username_penerima',
        'tipe'        => 'sppb_baru',
        'pesan'       => 'Unit X mengajukan SPPB baru: SPPB/2026/08/001',
        'url'         => '/logistik_non_medis/distribusisppb',
        'is_read'     => 0,
    ]);

[3] Controller memanggil fungsi _broadcastNotification($notification):
    - Menyiapkan data event JSON:
      { "type": "logistik_notification", "target": "username", "tipe": "sppb_baru",
        "pesan": "...", "url": "...", "id": 123 }
    - Menulis (APPEND) ke file: [BASE_DIR]/tmp/logistik_notifications.queue

[4] Workerman.php Timer (setiap 0.5 detik):
    - Cek apakah file queue ada dan ada konten baru (filesize > $queueOffset)
    - Buka file dengan flock(LOCK_SH) agar thread-safe
    - fseek ke $queueOffset (posisi terakhir yang sudah dibaca)
    - Baca baris baru dengan fgets() sampai EOF
    - Decode JSON setiap baris
    - Jika type == 'logistik_notification': broadcast ke semua koneksi WebSocket
    - Update $queueOffset ke posisi akhir
    - Release flock

[5] Browser (JavaScript) menerima event WebSocket:
    - Tampilkan toast notification
    - Update badge counter di bell icon
    - Jika user_target === username saat ini, tampilkan notif personal""")

doc.add_heading('3.3 Cara Menjalankan Workerman', level=2)
add_code(doc, """# 1. Install dependensi via Composer (dari root project)
composer install

# 2. Jalankan Workerman (buka CMD/Terminal terpisah)
php workerman.php start

# 3. Untuk mode daemon (background, Linux only):
php workerman.php start -d

# 4. Cek status server
php workerman.php status

# 5. Hentikan server
php workerman.php stop

# Konfigurasi Port: 3892 (bisa diubah di workerman.php baris ke-4)
# Pastikan port 3892 tidak diblokir firewall atau dipakai proses lain

# Folder yang harus ada dan bisa ditulis:
mkdir tmp
# Windows: klik kanan > Properties > Security > Allow write

# Cek koneksi WebSocket dari browser (Console DevTools):
# var ws = new WebSocket('ws://localhost:3892');
# ws.onmessage = function(e) { console.log(JSON.parse(e.data)); }""")

doc.add_heading('3.4 Trigger Notifikasi yang Sudah Diimplementasi', level=2)

add_table(doc, ['Kejadian', 'Tipe Notif', 'Penerima', 'Isi Pesan'],
[
    ['SPPB baru diajukan oleh unit', 'sppb_baru', 'Semua user dengan role admin & logistik', '"Unit [nama_unit] mengajukan SPPB baru: [no_sppb]"'],
    ['Status SPPB berubah menjadi "Siap Ambil"', 'siap_ambil', 'User yang berasal dari kode_unit pengaju SPPB', '"Permintaan Anda [no_sppb] sudah Siap Diambil di logistik"'],
], header_bg='D6F7E4')

doc.add_heading('3.5 Setting Konfigurasi Notifikasi', level=2)
doc.add_paragraph('Mode notifikasi disimpan di tabel settings mLITE dengan key berikut:')
add_code(doc, """Key   : logistik_non_medis.notif_realtime
Value : workerman   (mode saat ini)
        atau 'pusher' untuk mode Pusher/Ably (tidak aktif)

Cara mengubah via database:
UPDATE mlite_settings SET value = 'workerman'
WHERE name = 'logistik_non_medis.notif_realtime';""")

# SIMPAN
output_path = 'plugins/logistik_non_medis/DOKUMENTASI_PERUBAHAN_RSNS.docx'
doc.save(output_path)
print(f"DOCX berhasil dibuat: {output_path}")
