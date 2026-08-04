from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        run = h.runs[0] if h.runs else h.add_run(text)
        run.font.color.rgb = RGBColor(*color)
    return h

def add_table(doc, headers, rows, header_bg=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0] if hdr_cells[i].paragraphs[0].runs else hdr_cells[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        if header_bg:
            tc = hdr_cells[i]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), header_bg)
            tcPr.append(shd)
    # Data rows
    for ri, row_data in enumerate(rows):
        row_cells = table.rows[ri+1].cells
        for ci, cell_val in enumerate(row_data):
            row_cells[ci].text = str(cell_val)
            for para in row_cells[ci].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    return table

def add_code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    p.paragraph_format.left_indent = Inches(0.3)
    # Light gray background
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F5F5F5')
    pPr.append(shd)
    return p

# ========================
doc = Document()

# === PAGE MARGINS ===
section = doc.sections[0]
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

# === JUDUL ===
title = doc.add_heading('DOKUMENTASI PENYESUAIAN MODUL LOGISTIK NON MEDIS', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x15, 0x63, 0xBE)

sub = doc.add_paragraph('RSNS mLITE — Berdasarkan Commit & Perubahan Terbaru (Rifangga / Hanafi)')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(10)
sub.runs[0].italic = True

doc.add_paragraph(f'Digenerate: {datetime.datetime.now().strftime("%d %B %Y, %H:%M")}').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# ==========================================
# BAGIAN 1: PENYESUAIAN MENU & FITUR
# ==========================================
doc.add_heading('1. PENYESUAIAN MENU & FITUR ANTARMUKA', level=1)

doc.add_heading('1.1 Sistem Hak Akses (RBAC)', level=2)
doc.add_paragraph('Modul kini menggunakan sistem hak akses berbasis peran (Role-Based Access Control). Setiap role mendapat akses ke endpoint tertentu, disimpan di tabel role_permissions.')

add_table(doc, ['Role', 'Akses Endpoint', 'Keterangan'],
[
    ['admin', 'SEMUA endpoint', 'Akses penuh ke seluruh fitur dan konfigurasi'],
    ['logistik', 'manage, gudang*, distribusi*, perencanaan*, realisasi, masterinventaris, dll', 'Semua fitur operasional logistik'],
    ['aset', 'manage, aset*, masterinventaris, masterunit', 'Khusus pengelolaan aset dan inventaris'],
    ['gudang', 'manage, gudang*, distribusisppb, distribusiverifikasi, packing, serahterima, tracking, retur', 'Operasional gudang dan distribusi'],
    ['kepala_unit', 'manage, distribusisppb, distribusiverifikasi, distribusikuota', 'Persetujuan SPPB tingkat kepala unit'],
    ['kepala_sie', 'manage, distribusisppb, distribusiverifikasi, distribusikuota', 'Persetujuan SPPB tingkat kepala sie'],
    ['kepala_bidang', 'manage, distribusisppb, distribusiverifikasi, distribusikuota', 'Persetujuan SPPB tingkat kepala bidang'],
    ['unit', 'manage, distribusisppb, distribusikuota', 'Hanya mengajukan dan memantau SPPB unit sendiri'],
], header_bg='D6E4F7')

doc.add_heading('1.2 Menu Baru yang Ditambahkan', level=2)

add_table(doc, ['Nama Menu / Halaman', 'URL / Endpoint', 'File View', 'Keterangan'],
[
    ['Hak Akses Pengguna', '/logistik_non_medis/hakakses', 'hakakses.html + hakakses.form.html', 'Manajemen mapping user ke role dan unit'],
    ['Master Inventaris (Barang)', '/logistik_non_medis/masterinventaris', 'master.inventaris.html + master.inventaris.barang.display.html', 'Kelola master barang inventaris non-medis beserta hierarki kategori'],
    ['Rekanan/Vendor', '/logistik_non_medis/masterrekanan', 'master.rekanan.html', 'Master rekanan/vendor pengadaan'],
    ['Master Vendor', '/logistik_non_medis/mastervendor', 'master.vendor.html', 'Alternatif vendor evaluation'],
    ['Kategori Aset KIB', '/logistik_non_medis/masterkategoriaset', 'master_kategori_aset.html', 'Kelola kategori KIB (A=Tanah, B=Mesin, dst)'],
    ['Notifikasi', '/logistik_non_medis/notifikasi', 'notifikasi.html', 'Daftar notifikasi real-time per user'],
    ['Distribusi Non Rutin (SPPB Non Rutin)', '/logistik_non_medis/distribusinonrutin', 'distribusi.sppb.html (jenis_page=Non Rutin)', 'Alur pengajuan pengadaan non-rutin terpisah dari rutin'],
    ['Opname V2', '/logistik_non_medis/gudangopnamev2', 'gudang.opnamev2.html', 'Fitur opname stok versi baru yang lebih sederhana'],
    ['Realisasi Belanja', '/logistik_non_medis/realisasibelanja', 'pengadaan.realisasi.html', 'Monitoring realisasi belanja pengadaan'],
    ['Permintaan Non Rutin (Form)', '/logistik_non_medis/distribusinonrutinform', 'pengadaan.permintaan_nonrutin.html', 'Form pengajuan barang non-rutin oleh unit'],
], header_bg='D6F7E4')

doc.add_heading('1.3 Menu / Fitur yang Diubah', level=2)

add_table(doc, ['Halaman / Fitur', 'Perubahan', 'Detail'],
[
    ['Distribusi SPPB (Rutin)', 'Tab Aktif & Riwayat diperbaiki', 'Sekarang benar memfilter untuk semua role termasuk Admin & Logistik (sebelumnya hanya unit)'],
    ['Distribusi SPPB (Rutin)', 'Tombol "Import SPPB Mingguan" DIHAPUS', 'Fitur import mingguan dihilangkan dari antarmuka'],
    ['Dashboard Logistik', 'Widget counter diperluas', 'Counter baru: SPPB masuk minggu ini, menunggu verifikasi, dalam proses packing'],
    ['Form Registrasi Aset', 'Dropdown Kategori-Kelompok-Jenis ditambahkan', 'Sebelumnya tidak ada; kini terhubung ke tabel inventaris_kategori/kelompok/jenis'],
    ['List Registrasi Aset', 'Pagination diperbaiki & data tampil lengkap', 'Sebelumnya hanya tampil 5 data karena bug div di tbody; kini menggunakan response JSON'],
    ['Aset KIB', 'Kolom baru di tampilan: kib_jenis, nilai_buku, masa_manfaat', 'Mendukung format KIB A-F sesuai Permendagri'],
    ['Form Unit', 'Sinkronisasi PJ ke user_roles', 'Saat PJ unit diubah, data user_roles otomatis ikut diperbarui'],
    ['Halaman Manage (Dashboard)', 'Tampilan tab disesuaikan dengan role', 'Tab Distribusi, Pengadaan, Master tersembunyi jika tidak memiliki izin'],
    ['manage.html', 'Penambahan tab & badge counter baru', 'Badge notifikasi SPPB pending ditampilkan di tab Distribusi'],
], header_bg='F7EAD6')

doc.add_heading('1.4 Alur SPPB Non Rutin (Baru)', level=2)
doc.add_paragraph('Alur persetujuan SPPB Non Rutin berbeda dari Rutin:')
flow_steps = [
    ('Draft', 'Unit/Staff membuat draft permintaan non-rutin'),
    ('Diajukan', 'Unit/Staff mengirimkan permintaan'),
    ('Disetujui Ka. Sie', 'Kepala Sie menyetujui (level pertama untuk Non Rutin)'),
    ('Disetujui Kabid', 'Kepala Bidang menyetujui (level kedua)'),
    ('Proses Logistik', 'Tim Logistik memproses pengadaan'),
    ('Proses Pengadaan', 'Sedang dalam proses pengadaan (PO/tender)'),
    ('Siap Diserahkan', 'Barang sudah siap diserahkan ke unit'),
    ('Selesai / Diterima', 'Unit menerima barang dan SPPB selesai'),
    ('Ditolak', 'Ditolak di salah satu tahap (dicatat di kolom ditolak_pada_status)'),
]
table = doc.add_table(rows=len(flow_steps)+1, cols=2)
table.style = 'Table Grid'
table.rows[0].cells[0].text = 'Status'
table.rows[0].cells[1].text = 'Keterangan'
for i, (s, k) in enumerate(flow_steps):
    table.rows[i+1].cells[0].text = s
    table.rows[i+1].cells[1].text = k

doc.add_page_break()

# ==========================================
# BAGIAN 2: PENYESUAIAN TABEL DATABASE
# ==========================================
doc.add_heading('2. PENYESUAIAN TABEL DATABASE', level=1)
p = doc.add_paragraph('Semua tabel kustom modul Logistik Non Medis menggunakan awalan ')
run = p.add_run('rsns_custom_logistik_non_medis_')
run.bold = True
run.font.name = 'Courier New'
p.add_run('. Di bawah ini adalah ringkasan setiap perubahan kolom dan tabel yang dilakukan.')

# --- 2.1 Tabel sppb ---
doc.add_heading('2.1 Tabel: sppb', level=2)
doc.add_paragraph('Tabel utama permintaan barang (SPPB). Beberapa kolom baru ditambahkan dan kolom status dimodifikasi.')

add_table(doc, ['Jenis Perubahan', 'Kolom/Item', 'Detail', 'SQL / Keterangan'],
[
    ['ALTER — MODIFY', 'status (ENUM)', 'Diperluas dengan status baru', "MODIFY `status` ENUM('Draft','Diajukan','Disetujui Ka. Unit','Disetujui Ka. Sie','Disetujui Kabid','Disetujui Unit','Proses','Terverifikasi','Proses Logistik','Proses Pengadaan','Siap Ambil','Siap Diserahkan','Picking','Packing','Ready','Dikirim','Diterima','Selesai','Ditolak','Dibatalkan')"],
    ['ALTER — ADD', 'ditolak_pada_status', 'Menyimpan status saat SPPB ditolak', "ADD `ditolak_pada_status` varchar(100) DEFAULT NULL AFTER `alasan_penolakan`"],
    ['ALTER — ADD', 'jenis_permintaan', "Membedakan 'Rutin' dan 'Non Rutin'", "ADD `jenis_permintaan` varchar(50) NOT NULL DEFAULT 'Rutin' AFTER `kode_unit`"],
    ['ALTER — ADD', 'minggu_ke', 'Nomor minggu dalam tahun (untuk import mingguan)', "ADD `minggu_ke` int(11) DEFAULT NULL AFTER `tgl_sppb`"],
    ['ALTER — ADD (via commit terbaru)', 'user_approve_*', 'Kolom audit siapa yang menyetujui per level', "user_approve_ka_unit, user_approve_ka_sie, user_approve_ka_bidang, diambil_oleh, tgl_diambil"],
    ['DATA MIGRATION', 'status', "Rename status lama: 'Terverifikasi' → 'Proses Logistik'", "UPDATE sppb SET status = 'Proses Logistik' WHERE status = 'Terverifikasi'"],
    ['DATA MIGRATION', 'status', "Rename: 'Disetujui Unit' (Non Rutin) → 'Disetujui Kabid'", "UPDATE sppb SET status = 'Disetujui Kabid' WHERE status = 'Disetujui Unit' AND jenis_permintaan = 'Non Rutin'"],
], header_bg='D6E4F7')

# --- 2.2 Tabel aset ---
doc.add_heading('2.2 Tabel: aset (Registrasi & KIB)', level=2)
doc.add_paragraph('Kolom-kolom baru ditambahkan untuk mendukung fitur KIB (Kartu Inventaris Barang) yang lebih lengkap.')

add_table(doc, ['Jenis Perubahan', 'Kolom Baru', 'Tipe Data', 'Keterangan'],
[
    ['ALTER — ADD', 'nomor_inventaris', 'varchar(50)', 'Nomor inventaris aset manual (di luar kode AST otomatis)'],
    ['ALTER — ADD', 'merk_type', 'varchar(150)', 'Merk / tipe perangkat / aset'],
    ['ALTER — ADD', 'nomor_dokumen', 'varchar(200)', 'Nomor dokumen pengadaan aset'],
    ['ALTER — ADD', 'lokasi_fisik', 'varchar(150)', 'Lokasi fisik aset disimpan/dipasang'],
    ['ALTER — ADD', 'bahan', 'varchar(100)', 'Bahan/material aset'],
    ['ALTER — ADD', 'tahun_beli', 'smallint(4)', 'Tahun perolehan aset'],
    ['ALTER — ADD', 'satuan', 'varchar(50)', "Satuan aset (contoh: 'Buah', 'Unit')"],
    ['ALTER — ADD', 'jumlah', 'int(11) DEFAULT 1', 'Jumlah aset yang terdaftar'],
    ['ALTER — ADD', 'keterangan_inventaris', 'text', 'Catatan khusus tentang kondisi/lokasi'],
    ['ALTER — ADD KEY', 'INDEX nomor_inventaris', '-', 'Index baru untuk mempercepat pencarian nomor inventaris'],
    ['KIB kolom baru', 'kib_jenis', "ENUM('A','B','C','D','E','F')", 'Golongan KIB barang (A=Tanah, B=Mesin, dst)'],
    ['KIB kolom baru (aset)', 'kib_tgl_mulai, kib_tgl_rencana_selesai', 'date', 'Periode penggunaan aset dalam KIB'],
    ['KIB kolom baru (aset)', 'kib_nilai_buku, kib_akumulasi_penyusutan', 'decimal(15,2)', 'Nilai buku dan akumulasi penyusutan aset'],
], header_bg='D6F7E4')

# --- 2.3 Tabel inventaris_master ---
doc.add_heading('2.3 Tabel: inventaris_master', level=2)
doc.add_paragraph('Tabel master barang dan unit logistik non-medis diperluas untuk mendukung hierarki kategori.')

add_table(doc, ['Jenis Perubahan', 'Kolom', 'Tipe Data', 'Keterangan'],
[
    ['ALTER — ADD', 'kode_inventaris', 'char(3)', 'Kode inventaris singkat (untuk tabel UNIT)'],
    ['ALTER — MODIFY', 'kode', 'varchar(50)', 'Diperpanjang dari panjang sebelumnya'],
    ['ALTER — ADD', 'kode_kategori', 'char(1) DEFAULT ""', 'Kategori (1=Inventaris, 2=BHP, dsb)'],
    ['ALTER — ADD', 'nama_kelompok', 'varchar(150)', 'Nama kelompok barang (denormalisasi)'],
    ['ALTER — ADD', 'nama_jenis', 'varchar(150)', 'Nama jenis barang (denormalisasi)'],
    ['ALTER — ADD', 'kib_jenis', "ENUM('A'-'F')", 'Golongan KIB untuk master barang'],
    ['ALTER — DROP INDEX + ADD UNIQUE', 'jenis_kode', '-', 'Index unik diubah: (jenis_master, kode_kategori, kode)'],
], header_bg='F7EAD6')

# --- 2.4 Tabel Baru ---
doc.add_heading('2.4 Tabel Baru yang Dibuat', level=2)

add_table(doc, ['Nama Tabel', 'Fungsi', 'Kolom Utama'],
[
    ['notifikasi', 'Menyimpan notifikasi real-time per user', 'id, user_target, tipe, pesan, url, is_read, tgl_dibuat'],
    ['inventaris_kategori', 'Hierarki Kategori master barang', 'kode_kategori, nama_kategori, kib_default, umur_manfaat_default, kode_coa, status_aktif'],
    ['inventaris_kelompok', 'Hierarki Kelompok dalam kategori', 'kode_kategori, kode_kelompok, nama_kelompok'],
    ['inventaris_jenis', 'Hierarki Jenis dalam kelompok', 'kode_kategori, kode_kelompok, kode_jenis, nama_jenis'],
    ['role_permissions', 'Pemetaan hak akses per role', 'role, permissions (comma-separated endpoint slugs)'],
    ['user_roles', 'Pemetaan user ke role & unit', 'username, role, kode_unit'],
    ['kategori_aset', 'Kategori pengelompokan aset KIB', 'kode_kategori, nama_kategori, kib_default, umur_manfaat_default, kode_coa, status_aktif'],
    ['opname (v2)', 'Opname stok gudang versi 2', 'no_opname, tgl_opname, kode_gudang, status, dibuat_oleh'],
], header_bg='EAD6F7')

doc.add_paragraph()
doc.add_heading('2.5 Index Performa yang Ditambahkan', level=2)
doc.add_paragraph('File: performance_indexes.sql — Index baru ditambahkan untuk mempercepat query pada tabel dengan data besar:')
add_code_block(doc, """-- sppb
CREATE INDEX IF NOT EXISTS idx_sppb_status ON rsns_custom_logistik_non_medis_sppb (status);
CREATE INDEX IF NOT EXISTS idx_sppb_kode_unit ON rsns_custom_logistik_non_medis_sppb (kode_unit);
CREATE INDEX IF NOT EXISTS idx_sppb_tgl ON rsns_custom_logistik_non_medis_sppb (tgl_sppb);
-- aset
CREATE INDEX IF NOT EXISTS idx_aset_status ON rsns_custom_logistik_non_medis_aset (status);
CREATE INDEX IF NOT EXISTS idx_aset_kode_unit ON rsns_custom_logistik_non_medis_aset (kode_unit);""")

doc.add_page_break()

# ==========================================
# BAGIAN 3: PENGATURAN NOTIF WORKERMAN
# ==========================================
doc.add_heading('3. PENGATURAN NOTIFIKASI WORKERMAN', level=1)
doc.add_paragraph('Berikut adalah detail teknis perubahan pada sistem notifikasi real-time yang dilakukan berdasarkan commit terakhir.')

doc.add_heading('3.1 Gambaran Arsitektur Baru', level=2)
doc.add_paragraph('Sistem notifikasi sekarang menggunakan pola Queue File untuk menghubungkan PHP web server dengan Workerman WebSocket server, menggantikan pendekatan direct HTTP/cURL yang rentan error.')

add_code_block(doc, """Alur Notifikasi:
1. User melakukan aksi (misal: submit SPPB / status "Siap Ambil")
2. PHP controller memanggil $this->_broadcastNotification($data)
3. Fungsi tersebut MENULIS baris JSON ke: [BASE_DIR]/tmp/logistik_notifications.queue
4. Workerman WebSocket server (workerman.php) memiliki Timer yang berjalan setiap 0.5 detik
5. Timer membaca baris baru dari queue file menggunakan fseek() + flock()
6. Setiap baris JSON dengan type='logistik_notification' di-broadcast ke semua koneksi WebSocket aktif
7. Browser klien menerima event dan menampilkan toast/badge notifikasi""")

doc.add_heading('3.2 File workerman.php — Perubahan Teknis', level=2)

add_table(doc, ['Komponen', 'Sebelum', 'Sesudah'],
[
    ['Server WebSocket', 'Hanya mendengarkan koneksi dan onMessage dari browser', 'Ditambahkan Timer polling queue file setiap 0.5 detik'],
    ['Pengiriman Notifikasi dari PHP', 'Tidak ada mekanisme khusus (cURL ke port)', 'PHP menulis ke file queue; Workerman membaca file tersebut'],
    ['Broadcast', 'Hanya broadcast pesan yang diterima dari browser', 'Dua jalur: dari queue file (notif dari PHP) & dari browser (backward-compatible)'],
    ['Dependency Baru', '-', 'use Workerman\\Timer; ditambahkan'],
    ['Queue File', 'Tidak ada', '[BASE_DIR]/tmp/logistik_notifications.queue'],
    ['Offset tracking', 'Tidak ada', 'Variable $queueOffset memastikan baris yang sudah dibaca tidak dibaca ulang'],
], header_bg='D6E4F7')

doc.add_heading('3.3 Fungsi PHP Baru: _broadcastNotification()', level=2)
doc.add_paragraph('Fungsi private baru di Admin.php yang menangani pengiriman notifikasi dari backend:')
add_code_block(doc, """private function _broadcastNotification($notification) {
    $event = [
        'type'   => 'logistik_notification',
        'target' => $notification['user_target'] ?? '',
        'tipe'   => $notification['tipe'] ?? '',
        'pesan'  => $notification['pesan'] ?? '',
        'url'    => $notification['url'] ?? '',
        'id'     => $notification['id'] ?? null,
    ];
    $queue = BASE_DIR . '/tmp/logistik_notifications.queue';
    // Cek mode (workerman atau pusher) dari settings
    $mode = $this->settings->get('logistik_non_medis.notif_realtime');
    // Tulis event ke queue file untuk dibaca Workerman
    file_put_contents($queue, json_encode($event) . "\\n", FILE_APPEND | LOCK_EX);
}""")

doc.add_heading('3.4 Tabel Notifikasi (Baru)', level=2)
doc.add_paragraph('Notifikasi juga disimpan ke database untuk riwayat dan halaman notifikasi:')
add_code_block(doc, """CREATE TABLE IF NOT EXISTS `rsns_custom_logistik_non_medis_notifikasi` (
    `id`          INT AUTO_INCREMENT PRIMARY KEY,
    `user_target` VARCHAR(100) NOT NULL,   -- username penerima notifikasi
    `tipe`        VARCHAR(50)  NOT NULL,   -- 'sppb_baru', 'siap_ambil', dll
    `pesan`       TEXT         NOT NULL,   -- isi pesan notifikasi
    `url`         VARCHAR(500) DEFAULT NULL, -- link tujuan saat notif diklik
    `is_read`     TINYINT(1)   NOT NULL DEFAULT 0,
    `tgl_dibuat`  DATETIME     NOT NULL DEFAULT CURRENT_TIMESTAMP,
    INDEX (`user_target`),
    INDEX (`is_read`)
);""")

doc.add_heading('3.5 Trigger Notifikasi', level=2)
doc.add_paragraph('Notifikasi dikirim secara otomatis pada dua kejadian:')

add_table(doc, ['Kejadian / Trigger', 'Tipe Notifikasi', 'Penerima', 'Isi Pesan'],
[
    ['SPPB baru diajukan', 'sppb_baru', 'Semua user role admin & logistik', 'Unit X mengajukan SPPB baru (No. SPPB/...)'],
    ['Status SPPB → "Siap Ambil"', 'siap_ambil', 'User dari kode_unit yang mengajukan', 'Permintaan Anda (No. SPPB/...) sudah Siap Diambil'],
], header_bg='D6F7E4')

doc.add_heading('3.6 Cara Menjalankan Workerman', level=2)
add_code_block(doc, """# Install dependensi dulu (pastikan composer.json sudah benar)
composer install

# Jalankan Workerman di background (Windows: gunakan CMD terpisah)
php workerman.php start

# Cek status
php workerman.php status

# Hentikan
php workerman.php stop

# Konfigurasi port WebSocket: 3892 (default, bisa diubah di workerman.php)
# Pastikan folder tmp/ bisa ditulis oleh web server:
# mkdir tmp   (jika belum ada)
# chmod 775 tmp   (Linux)""")

doc.add_heading('3.7 Konfigurasi Pengaturan Notif di mLITE', level=2)
doc.add_paragraph('Setting mode notifikasi disimpan di tabel settings mLITE dengan key:')
add_code_block(doc, """Key   : logistik_non_medis.notif_realtime
Value : workerman  (atau 'pusher' untuk mode Pusher/Ably jika tersedia)""")

doc.add_paragraph('Setting ini bisa diubah dari halaman Pengaturan Logistik Non Medis (menu Admin).')

doc.add_heading('3.8 Struktur Folder yang Dibutuhkan', level=2)
add_code_block(doc, """[root mlite_rsns]/
├── workerman.php          <- Server WebSocket (DIUBAH)
├── tmp/
│   └── logistik_notifications.queue   <- Queue file notifikasi (auto-create)
├── vendor/
│   └── workerman/          <- Library Workerman (via composer)
└── plugins/logistik_non_medis/
    └── view/admin/
        └── notifikasi.html <- Halaman daftar notifikasi (BARU)""")

doc.add_paragraph()
doc.add_paragraph('Note: Queue file akan dibuat otomatis oleh PHP saat notifikasi pertama dikirim. Workerman akan terus memantaunya selama server berjalan.').italic = True

# === SIMPAN ===
output_path = 'plugins/logistik_non_medis/DOKUMENTASI_PENYESUAIAN_v2.docx'
doc.save(output_path)
print(f"DOCX berhasil dibuat: {output_path}")
