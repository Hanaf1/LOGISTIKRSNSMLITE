import os
import subprocess
import sys

def install_and_import(package):
    try:
        __import__(package)
    except ImportError:
        print(f"Installing {package}...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", package])

install_and_import('docx')
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# Set default font to Times New Roman and Black for Normal
style_normal = doc.styles['Normal']
font_normal = style_normal.font
font_normal.name = 'Times New Roman'
font_normal.size = Pt(12)
font_normal.color.rgb = RGBColor(0, 0, 0)

# Apply to all heading styles to ensure TNR and Black
for i in range(1, 10):
    try:
        style_heading = doc.styles[f'Heading {i}']
        font_heading = style_heading.font
        font_heading.name = 'Times New Roman'
        font_heading.color.rgb = RGBColor(0, 0, 0)
    except KeyError:
        pass
        
try:
    style_title = doc.styles['Title']
    font_title = style_title.font
    font_title.name = 'Times New Roman'
    font_title.color.rgb = RGBColor(0, 0, 0)
except KeyError:
    pass

# Title
title = doc.add_heading('LAPORAN KHUSUS KARYAWAN', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('USULAN DIGITALISASI PROSES LOGISTIK & GUDANG RS', 1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Instansi: RSU Nurusyifa', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Disusun Oleh: Admin Gudang\n', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER

# 1. PENDAHULUAN
doc.add_heading('1. PENDAHULUAN', level=2)
doc.add_heading('1.1. Latar Belakang', level=3)
doc.add_paragraph('Seiring dengan perkembangan teknologi informasi dan meningkatnya beban pelayanan medis maupun non-medis di RSU Nurusyifa, pengelolaan logistik dan gudang secara manual atau konvensional seringkali memunculkan berbagai kendala. Masalah yang sering terjadi meliputi keterlambatan dalam pemrosesan permintaan barang (rutin dan insidental), ketidaksesuaian antara pencatatan stok di kartu stok fisik dengan ketersediaan barang sebenarnya, hingga proses pengadaan yang memakan waktu lama akibat rantai birokrasi persetujuan (approval) yang berbasis kertas (paper-based).\n\nOleh karena itu, dibutuhkan sebuah inovasi berupa digitalisasi sistem informasi manajemen logistik dan gudang. Sistem ini dirancang untuk mengotomatisasi, mengintegrasikan, dan memonitor seluruh rantai pasok (supply chain) rumah sakit secara real-time, efektif, dan efisien.', style='Normal')

doc.add_heading('1.2. Tujuan dan Sasaran', level=3)
doc.add_paragraph('Tujuan utama dari usulan digitalisasi ini meliputi:', style='Normal')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Mempercepat proses permintaan barang (amprah) dari seluruh unit/ruangan pelayanan ke gudang pusat.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Menjaga akurasi dan sinkronisasi data stok barang (Inventory Control) secara real-time.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Menyediakan sistem pelacakan (tracking) status pesanan dan pengadaan secara transparan.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Mewujudkan lingkungan kerja yang paperless dan meminimalisir kesalahan manusia (human error) akibat rekapitulasi manual.')

# 2. ANALISIS KONDISI DAN SOLUSI
doc.add_heading('2. ANALISIS KONDISI DAN SOLUSI', level=2)
doc.add_paragraph('Beberapa kendala yang sering dijumpai pada sistem manual antara lain hilangnya form permintaan, staf ruangan harus bolak-balik ke gudang hanya untuk mengecek ketersediaan barang, dan kesulitan dalam melacak riwayat penggunaan barang per ruangan. Solusi digital ini akan mengubah alur kerja dari reaktif menjadi proaktif dengan dukungan data yang valid.', style='Normal')

# 3. DETAIL FITUR SISTEM DAN ALUR OPERASIONAL
doc.add_heading('3. DETAIL FITUR SISTEM DAN ALUR OPERASIONAL', level=2)
doc.add_paragraph('Sistem logistik yang baru akan dibangun mencakup beberapa modul utama yang saling terintegrasi:', style='Normal')

# Modul Master Data
doc.add_heading('3.1. Modul Manajemen Master Data', level=3)
doc.add_paragraph('Modul ini merupakan pondasi awal sistem yang mengatur kelengkapan data referensi:', style='Normal')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Master Barang & Kategori: ').bold = True
p.add_run('Pengelompokan barang logistik non-medis, ATK, alkes, dll, dilengkapi dengan satuan ukur (box, pcs, botol).')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Master Rekanan/Supplier: ').bold = True
p.add_run('Penyimpanan profil vendor untuk mempermudah penerbitan PO.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Pengaturan Buffer Stock (Stok Minimum): ').bold = True
p.add_run('Fitur untuk menetapkan batas aman stok. Jika menyentuh batas ini, sistem akan otomatis memberi peringatan.')

# Modul Dashboard
doc.add_heading('3.2. Modul Dashboard & Monitoring Gudang', level=3)
doc.add_paragraph('Berfungsi sebagai pusat kendali bagi Admin Gudang.', style='Normal')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Notifikasi Defecta (Stok Kritis): ').bold = True
p.add_run('Tampilan langsung pada layar untuk item yang harus segera di-restock.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Ringkasan Status Permintaan: ').bold = True
p.add_run('Menampilkan indikator status (Berapa permintaan baru, berapa yang sedang disiapkan, dan yang selesai).')
try:
    doc.add_picture(r'C:\Users\ASUS\.gemini\antigravity-ide\brain\4e43103c-513c-43c9-83bf-d5fea4699698\dashboard_logistik_1785577433907.png', width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_caption = doc.add_paragraph('Gambar 1: Dashboard Monitoring Gudang', style='Normal')
    p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
except Exception as e:
    pass

# Modul Permintaan Unit
doc.add_heading('3.3. Modul Permintaan Barang (Unit Pelayanan)', level=3)
doc.add_paragraph('Setiap unit akan memiliki akses ke E-Katalog internal RSU Nurusyifa.', style='Normal')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Sistem E-Katalog & Cart: ').bold = True
p.add_run('Mirip dengan aplikasi e-commerce, ruangan dapat mencari barang, melihat sisa stok yang ada di gudang, dan memasukkan jumlah yang diminta ke dalam "Keranjang".')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Validasi Stok Otomatis: ').bold = True
p.add_run('Sistem menolak permintaan jika jumlah yang diminta melebihi sisa stok di gudang, kecuali untuk permintaan insidental/darurat.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Live Tracking: ').bold = True
p.add_run('Ruangan tidak perlu menelepon gudang karena status pengajuan bisa dilihat langsung (Draft -> Menunggu Persetujuan -> Disiapkan Gudang -> Terkirim).')
try:
    doc.add_picture(r'C:\Users\ASUS\.gemini\antigravity-ide\brain\4e43103c-513c-43c9-83bf-d5fea4699698\form_permintaan_1785577447077.png', width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_caption = doc.add_paragraph('Gambar 2: Form E-Katalog Permintaan Barang Ruangan', style='Normal')
    p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
except Exception as e:
    pass

# Modul Distribusi
doc.add_heading('3.4. Modul Mutasi dan Distribusi Internal', level=3)
doc.add_paragraph('Memfasilitasi alur pengeluaran barang dari Gudang Utama ke Unit/Sub-Gudang.', style='Normal')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Cetak Surat Jalan Internal: ').bold = True
p.add_run('Admin gudang mencetak struk/dokumen pengiriman sebagai bukti fisik saat mendistribusikan barang.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Digital Handover (Serah Terima): ').bold = True
p.add_run('Saat barang tiba di ruangan, pihak ruangan menekan tombol "Terima Barang" di sistem, barulah stok di Gudang Pusat resmi terpotong dan mutasi stok tercatat.')

# Modul Pengadaan
doc.add_heading('3.5. Modul Pengadaan (Procurement) dan Approval Berjenjang', level=3)
doc.add_paragraph('Untuk memastikan tata kelola keuangan yang baik dan proses pengadaan yang efisien.', style='Normal')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Auto-Purchase Request (PR): ').bold = True
p.add_run('Admin gudang dapat meng-generate PR secara instan berdasarkan data Defecta (stok kurang).')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Approval Berjenjang (Paperless): ').bold = True
p.add_run('PR diteruskan secara hierarkis (misal: Kepala Logistik -> Direktur Keuangan). Approval dilakukan cukup dengan sekali klik.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Penerbitan PO: ').bold = True
p.add_run('Setelah disetujui, sistem membuat Purchase Order (PO) dalam format PDF yang siap dikirim langsung melalui email ke Supplier.')
try:
    doc.add_picture(r'C:\Users\ASUS\.gemini\antigravity-ide\brain\4e43103c-513c-43c9-83bf-d5fea4699698\pengadaan_approval_1785577462216.png', width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_caption = doc.add_paragraph('Gambar 3: Halaman Approval dan Purchase Order', style='Normal')
    p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
except Exception as e:
    pass

# Modul Penerimaan
doc.add_heading('3.6. Modul Penerimaan Barang (Goods Receipt)', level=3)
p = doc.add_paragraph(style='List Bullet')
p.add_run('Pencocokan Faktur: ').bold = True
p.add_run('Mencocokkan fisik barang yang dikirim oleh vendor dengan PO rumah sakit.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Manajemen Tanggal Kedaluwarsa: ').bold = True
p.add_run('Input batch number dan expired date untuk menerapkan sistem FEFO (First Expired First Out).')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Mutasi Masuk: ').bold = True
p.add_run('Stok gudang utama otomatis bertambah begitu penerimaan dikonfirmasi di sistem.')

# 4. KESIMPULAN & TINDAK LANJUT
doc.add_heading('4. KESIMPULAN & TINDAK LANJUT', level=2)
doc.add_paragraph('Digitalisasi ini bukan sekadar mengubah dokumen kertas menjadi file komputer, melainkan mengubah proses bisnis (Business Process Reengineering) menjadi jauh lebih cepat, terkontrol, dan transparan. Langkah selanjutnya yang diusulkan adalah tahap persiapan infrastruktur, integrasi dengan Modul Logistik Non-Medis mLite RSNS, dan penetapan jadwal uji coba secara bertahap.', style='Normal')

# Save doc
output_path = os.path.join(os.path.dirname(__file__), 'Laporan_Khusus_Karyawan_Sangat_Detail_TNR.docx')
doc.save(output_path)
print(f"Berhasil membuat dokumen DOCX di: {output_path}")
