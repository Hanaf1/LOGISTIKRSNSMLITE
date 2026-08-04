import os
import subprocess
import sys

def install_and_import(package):
    try:
        __import__(package)
    except ImportError:
        print(f"Installing {package}...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", package])

install_and_import('docx')
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('LAPORAN KHUSUS KARYAWAN', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('USULAN DIGITALISASI PROSES LOGISTIK & GUDANG (DETAIL SISTEM)', 1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('RSU Nurusyifa', style='Intense Quote')
doc.add_paragraph('Disusun Oleh: Admin Gudang\n')

# Section 1
doc.add_heading('1. PENDAHULUAN', level=2)
doc.add_heading('1.1. Latar Belakang', level=3)
doc.add_paragraph('Seiring dengan perkembangan teknologi dan meningkatnya kebutuhan pelayanan kesehatan di RSU Nurusyifa, pengelolaan gudang secara manual dirasa kurang efektif. Sering terjadi keterlambatan dalam pemrosesan permintaan barang (rutin maupun non-rutin), ketidaksesuaian stok fisik dengan catatan, serta proses pengadaan yang memakan waktu lama. Oleh karena itu, diperlukan suatu sistem digitalisasi untuk mengotomatisasi dan memonitor seluruh proses logistik gudang.')

doc.add_heading('1.2. Tujuan', level=3)
p = doc.add_paragraph(style='List Bullet')
p.add_run('Mempercepat proses permintaan barang dari tiap unit/ruangan.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Memastikan keakuratan data stok barang secara real-time.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Mempermudah pelacakan (tracking) status permintaan dan pengadaan barang.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Mengurangi penggunaan kertas (paperless) dan meminimalisir human error.')

# Section 2
doc.add_heading('2. DETAIL FITUR SISTEM LOGISTIK & GUDANG', level=2)
doc.add_paragraph('Sistem logistik yang baru akan dirancang agar mudah digunakan (user-friendly) dengan alur yang jelas. Berikut adalah penjabaran detail fitur pada sistem:')

# Fitur 1
doc.add_heading('2.1. Dashboard Utama (Gudang)', level=3)
doc.add_paragraph('Dashboard ini akan menjadi halaman pertama yang dilihat oleh Admin Gudang. Fitur utamanya meliputi:')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Alert Stok Kritis (Defecta):').bold = True
p.add_run(' Sistem otomatis memberi peringatan jika stok barang mencapai batas minimum.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Ringkasan Permintaan:').bold = True
p.add_run(' Menampilkan jumlah permintaan dari unit yang menunggu diproses (Pending, On-Progress, Selesai).')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Grafik Pemakaian:').bold = True
p.add_run(' Memvisualisasikan barang mana saja yang pergerakannya paling cepat (Fast Moving).')
try:
    doc.add_picture(r'C:\Users\ASUS\.gemini\antigravity-ide\brain\4e43103c-513c-43c9-83bf-d5fea4699698\dashboard_logistik_1785577433907.png', width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_caption = doc.add_paragraph('Gambar 1: Mockup Desain Dashboard Logistik Gudang', style='Caption')
    p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
except Exception as e:
    doc.add_paragraph(f'(Gambar tidak dapat dimuat: {e})')

# Fitur 2
doc.add_heading('2.2. Modul Permintaan Barang (Unit ke Gudang)', level=3)
doc.add_paragraph('Setiap unit ruangan akan diberikan akses untuk melakukan permintaan (baik rutin maupun non-rutin) secara digital. Fitur utamanya meliputi:')
p = doc.add_paragraph(style='List Bullet')
p.add_run('E-Katalog Barang:').bold = True
p.add_run(' Unit dapat mencari barang yang dibutuhkan lengkap dengan informasi stok yang masih tersedia di gudang utama.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Form Input Keranjang (Cart):').bold = True
p.add_run(' Memudahkan unit memilih banyak barang sekaligus seperti berbelanja online.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Keterangan Urgensi (Insidental):').bold = True
p.add_run(' Field khusus untuk mencantumkan alasan jika ada permintaan mendesak / non-rutin.')
try:
    doc.add_picture(r'C:\Users\ASUS\.gemini\antigravity-ide\brain\4e43103c-513c-43c9-83bf-d5fea4699698\form_permintaan_1785577447077.png', width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_caption = doc.add_paragraph('Gambar 2: Mockup Form Permintaan Barang dari Unit', style='Caption')
    p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
except Exception as e:
    doc.add_paragraph(f'(Gambar tidak dapat dimuat: {e})')

# Fitur 3
doc.add_heading('2.3. Modul Approval dan Pengadaan (Purchase Order)', level=3)
doc.add_paragraph('Sistem ini memastikan tidak ada barang keluar atau dipesan ke supplier tanpa otorisasi yang sah. Fitur utamanya meliputi:')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Approval Berjenjang:').bold = True
p.add_run(' Pihak berwenang (Manajer/Direktur) cukup klik "Setuju" atau "Tolak" dari perangkat mereka.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Auto-Generate PR/PO:').bold = True
p.add_run(' Jika stok habis dan pengajuan disetujui, sistem akan otomatis mencetak Surat Pemesanan (Purchase Order) yang bisa langsung dikirim via sistem atau email ke supplier.')
try:
    doc.add_picture(r'C:\Users\ASUS\.gemini\antigravity-ide\brain\4e43103c-513c-43c9-83bf-d5fea4699698\pengadaan_approval_1785577462216.png', width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_caption = doc.add_paragraph('Gambar 3: Mockup Halaman Approval dan PO', style='Caption')
    p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
except Exception as e:
    doc.add_paragraph(f'(Gambar tidak dapat dimuat: {e})')

# Section 3
doc.add_heading('3. ALUR OPERASIONAL HARIAN (SOP BARU)', level=2)
doc.add_paragraph('Berikut adalah alur operasional singkat jika sistem di atas diterapkan:')
p = doc.add_paragraph(style='List Number')
p.add_run('Permintaan:').bold = True
p.add_run(' Ruangan membuka sistem -> Cari barang di katalog -> Input jumlah -> Submit.')
p = doc.add_paragraph(style='List Number')
p.add_run('Validasi & Distribusi:').bold = True
p.add_run(' Gudang mengecek dashboard -> Menyetujui permintaan (jika stok ada) -> Menyiapkan fisik barang -> Ruangan mengambil/menerima -> Klik "Barang Diterima" di sistem -> Stok terpotong.')
p = doc.add_paragraph(style='List Number')
p.add_run('Pengadaan (Jika Stok Habis):').bold = True
p.add_run(' Gudang membuat PR -> Manajemen Approve -> Sistem membuat PO -> Barang datang -> Gudang input Penerimaan Barang -> Stok bertambah.')

# Section 4
doc.add_heading('4. KESIMPULAN', level=2)
doc.add_paragraph('Dengan pendetailan fitur ini (Dashboard, E-Katalog Permintaan, dan Approval Pengadaan Otomatis), diharapkan digitalisasi logistik RSU Nurusyifa dapat berjalan secara komprehensif, transparan, dan mempercepat layanan.')

# Save doc
output_path = os.path.join(os.path.dirname(__file__), 'Laporan_Khusus_Karyawan_Digitalisasi_Detail.docx')
doc.save(output_path)
print(f"Berhasil membuat dokumen DOCX di: {output_path}")
