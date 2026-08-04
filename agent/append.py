
from docx import Document
import sys

def main():
    doc_path = r'c:\laragon\www\mlite_rsns\agent\Laporan_Khusus_Karyawan_Sangat_Detail_TNR.docx'
    doc = Document(doc_path)

    sections = [
        ("4.1 Penggunaan oleh Unit/Ruangan", [
            "Petugas unit masuk ke sistem.",
            "Memilih barang melalui katalog.",
            "Mengisi jumlah permintaan.",
            "Mengirim permintaan rutin atau insidental.",
            "Memantau status permintaan.",
            "Melakukan konfirmasi penerimaan barang."
        ]),
        ("4.2 Penggunaan oleh Admin Gudang", [
            "Melihat permintaan yang masuk.",
            "Memeriksa stok barang.",
            "Menyetujui atau menyesuaikan jumlah permintaan.",
            "Menyiapkan dan mendistribusikan barang.",
            "Mencatat mutasi stok.",
            "Membuat rencana pengadaan ketika stok minimum tercapai."
        ]),
        ("4.3 Penggunaan oleh Pejabat Approval", [
            "Menerima notifikasi pengajuan pengadaan.",
            "Memeriksa rincian barang dan estimasi biaya.",
            "Menyetujui atau menolak pengajuan.",
            "Memantau penerbitan Purchase Order."
        ])
    ]

    for heading, bullets in sections:
        doc.add_heading(heading, level=2)
        for bullet in bullets:
            doc.add_paragraph(bullet, style='List Bullet')

    doc.save(doc_path)
    print("Successfully appended to document.")

if __name__ == "__main__":
    main()
