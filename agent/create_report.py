import os
import subprocess
import sys

def install_and_import(package):
    try:
        __import__(package)
    except ImportError:
        print(f"Installing {package}...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", package])

install_and_import('docx')
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('LAPORAN KHUSUS KARYAWAN', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('USULAN DIGITALISASI PROSES LOGISTIK & GUDANG', 1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('RSU Nurusyifa', style='Intense Quote')
doc.add_paragraph('Disusun Oleh: Admin Gudang\n')

# Section 1
doc.add_heading('1. PENDAHULUAN', level=2)
doc.add_heading('1.1. Latar Belakang', level=3)
doc.add_paragraph('Seiring dengan perkembangan teknologi dan meningkatnya kebutuhan pelayanan kesehatan di RSU Nurusyifa, pengelolaan gudang secara manual dirasa kurang efektif. Sering terjadi keterlambatan dalam pemrosesan permintaan barang (rutin maupun non-rutin), ketidaksesuaian stok fisik dengan catatan, serta proses pengadaan yang memakan waktu lama. Oleh karena itu, diperlukan suatu sistem digitalisasi untuk mengotomatisasi dan memonitor seluruh proses logistik gudang.')

doc.add_heading('1.2. Tujuan', level=3)
p = doc.add_paragraph(style='List Bullet')
p.add_run('Mempercepat proses permintaan barang dari tiap unit/ruangan.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Memastikan keakuratan data stok barang secara real-time.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Mempermudah pelacakan (tracking) status permintaan dan pengadaan barang.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Mengurangi penggunaan kertas (paperless) dan meminimalisir human error.')

# Section 2
doc.add_heading('2. PANDUAN ALUR DIGITALISASI PROSES GUDANG', level=2)
doc.add_paragraph('Berikut adalah panduan alur baru yang diusulkan dalam sistem digitalisasi logistik RSU Nurusyifa:')

doc.add_heading('2.1. Proses Permintaan Rutin (Unit/Ruangan ke Gudang)', level=3)
p = doc.add_paragraph(style='List Number')
p.add_run('Input Permintaan:').bold = True
p.add_run(' User/Kepala Ruangan melakukan input permintaan barang rutin melalui sistem (Aplikasi mLite/Sistem Informasi RS) dengan memilih item yang tersedia beserta jumlahnya.')
p = doc.add_paragraph(style='List Number')
p.add_run('Validasi & Verifikasi:').bold = True
p.add_run(' Admin Gudang menerima notifikasi permintaan dan melakukan verifikasi ketersediaan stok di sistem.')
p = doc.add_paragraph(style='List Number')
p.add_run('Approval:').bold = True
p.add_run(' Sistem akan meneruskan permintaan ke pihak berwenang (Manajer Logistik/Direktur) untuk persetujuan (approval) secara digital jika diperlukan.')
p = doc.add_paragraph(style='List Number')
p.add_run('Distribusi:').bold = True
p.add_run(' Setelah disetujui, Admin Gudang menyiapkan barang dan mengubah status di sistem menjadi "Siap Diambil" atau "Sedang Didistribusikan".')
p = doc.add_paragraph(style='List Number')
p.add_run('Serah Terima:').bold = True
p.add_run(' Pihak ruangan menerima barang dan melakukan konfirmasi penerimaan di sistem. Stok otomatis berkurang.')

doc.add_heading('2.2. Proses Permintaan Non-Rutin (Insidental)', level=3)
p = doc.add_paragraph(style='List Number')
p.add_run('Pengajuan:').bold = True
p.add_run(' Ruangan mengisi form permintaan non-rutin di sistem beserta alasan kebutuhan (urgensi).')
p = doc.add_paragraph(style='List Number')
p.add_run('Review Admin Gudang:').bold = True
p.add_run(' Admin mengecek apakah barang bisa dipenuhi dari stok atau perlu pengadaan baru.')
p = doc.add_paragraph(style='List Number')
p.add_run('Eskalasi Pengadaan:').bold = True
p.add_run(' Jika barang tidak tersedia, sistem otomatis membuat draft pengajuan pengadaan (Purchase Request).')

doc.add_heading('2.3. Proses Pengadaan Barang (Procurement)', level=3)
p = doc.add_paragraph(style='List Number')
p.add_run('Purchase Request (PR):').bold = True
p.add_run(' Berdasarkan rekap permintaan atau peringatan stok minimum (reorder point), Admin Gudang men-generate PR di sistem.')
p = doc.add_paragraph(style='List Number')
p.add_run('Approval PR & PO:').bold = True
p.add_run(' PR disetujui oleh Manajemen, dilanjutkan dengan pembuatan Purchase Order (PO) digital yang dikirim ke Supplier via email/sistem.')
p = doc.add_paragraph(style='List Number')
p.add_run('Penerimaan Barang:').bold = True
p.add_run(' Saat barang datang, Admin Gudang melakukan pengecekan fisik vs PO. Jika sesuai, dilakukan Penerimaan Barang di sistem (stok otomatis bertambah).')

# Section 3
doc.add_heading('3. IMPLEMENTASI DAN EVALUASI', level=2)
doc.add_heading('3.1. Kebutuhan Sistem', level=3)
doc.add_paragraph('Modul Logistik Non-Medis pada sistem mLite RSNS yang terintegrasi.', style='List Bullet')
doc.add_paragraph('Perangkat keras (Komputer/Tablet di tiap ruangan untuk input permintaan).', style='List Bullet')
doc.add_paragraph('Jaringan intranet/internet yang stabil.', style='List Bullet')

doc.add_heading('3.2. Tahapan Evaluasi', level=3)
doc.add_paragraph('Bulan 1: Sosialisasi SOP baru dan uji coba (pilot project) di beberapa ruangan.', style='List Bullet')
doc.add_paragraph('Bulan 2: Evaluasi kendala sistem dan user error, serta perbaikan bug.', style='List Bullet')
doc.add_paragraph('Bulan 3: Implementasi penuh di seluruh lingkungan RSU Nurusyifa.', style='List Bullet')

# Section 4
doc.add_heading('4. PENUTUP', level=2)
doc.add_paragraph('Dengan adanya digitalisasi proses permintaan dan pengadaan di gudang RSU Nurusyifa, diharapkan tercipta efisiensi waktu, transparansi alur barang, dan akurasi data inventaris yang lebih baik. Panduan ini dapat menjadi acuan awal dalam pengembangan sistem logistik ke depannya.')

# Save doc
output_path = os.path.join(os.path.dirname(__file__), 'Laporan_Khusus_Karyawan_Digitalisasi.docx')
doc.save(output_path)
print(f"Berhasil membuat dokumen DOCX di: {output_path}")
