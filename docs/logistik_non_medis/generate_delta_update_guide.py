from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from generate_deployment_guide import (
    BLUE,
    GREEN,
    LIGHT_BLUE,
    NAVY,
    ORANGE,
    RED,
    TEXT,
    WHITE,
    add_bullet,
    add_callout,
    add_checklist,
    add_code,
    add_number,
    add_table,
    add_title,
    configure,
    margins,
    shade,
)


OUTPUT_DIR = Path(__file__).resolve().parent
OUTPUT_FILE = OUTPUT_DIR / "05_Delta_Update_Rifangga_ke_Sistem_RSNS_Terbaru.docx"
DOC_DATE = "02 Agustus 2026"


def cover(doc):
    p = doc.add_paragraph("DOKUMEN DELTA TEKNIS", style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\n")
    title = doc.add_paragraph("PERBEDAAN DAN PENYESUAIAN UPDATE", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    product = doc.add_paragraph("RIFANGGA TERBARU -> SISTEM RSNS SEKARANG")
    product.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = product.runs[0]
    run.font.name = "Arial"
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(BLUE)

    band = doc.add_table(rows=1, cols=1)
    band.alignment = WD_TABLE_ALIGNMENT.CENTER
    band.autofit = False
    band.columns[0].width = Cm(15.8)
    cell = band.cell(0, 0)
    shade(cell, NAVY)
    margins(cell, 260, 240, 260, 240)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Peta file, struktur SQL, Workerman, Pusher, dan urutan penerapan")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor.from_string(WHITE)

    doc.add_paragraph("\n")
    add_table(
        doc,
        ["Acuan", "Nilai"],
        [
            ("Upstream Rifangga", "rifangga/main @ 027fcfd - 24 Juli 2026"),
            ("Sistem RSNS", "branch fitur-hak-akses-style-lama @ 39a5024 + working tree terbaru"),
            ("Tanggal analisis", DOC_DATE),
            ("Tujuan", "Memudahkan teknisi menerapkan sistem RSNS tanpa kehilangan data dan penyesuaian lokal"),
        ],
        [4.7, 12.0],
    )
    p = doc.add_paragraph("Dokumen ini hanya menjelaskan perbedaan dan penyesuaian teknis. Prosedur backup/deploy umum berada pada dokumen modul deployment terpisah.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()


def build_document():
    doc = Document()
    configure(doc)
    header = doc.sections[0].header.paragraphs[0]
    header.text = "DOKUMEN DELTA | RIFANGGA TERBARU VS SISTEM RSNS"
    header.runs[0].font.name = "Arial"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.bold = True
    header.runs[0].font.color.rgb = RGBColor.from_string(NAVY)
    cover(doc)

    add_title(doc, "Ringkasan Delta")
    add_callout(
        doc,
        "Upstream sudah terkandung",
        "Commit rifangga/main 027fcfd sudah menjadi ancestor sistem RSNS. Karena itu teknisi tidak perlu mengulang merge commit tersebut. Paket update harus dibuat dari sistem RSNS yang sudah dirapikan dan dikomit.",
        "success",
    )
    add_table(
        doc,
        ["Ukuran perubahan", "Hasil analisis"],
        [
            ("File berbeda dari Rifangga", "157 file"),
            ("Perubahan baris", "+17.931 / -2.782"),
            ("Method baru di Admin.php", "52 method"),
            ("Tabel runtime baru di Admin.php", "5 tabel"),
            ("Perubahan lokal belum dikomit", "27 file termodifikasi + 10 item untracked saat snapshot"),
        ],
        [6.4, 10.3],
    )
    add_callout(
        doc,
        "Makna angka 157 file",
        "Angka ini mencakup kode produksi, view, SQL, dokumen, file debug, CSV/XLSX, dan alat bantu. Jangan menyalin semuanya secara buta ke server. Gunakan klasifikasi pada Bab 3.",
        "warning",
    )

    add_title(doc, "1. Perbedaan Fungsional Utama")
    add_table(
        doc,
        ["Area", "Rifangga 027fcfd", "Sistem RSNS sekarang"],
        [
            ("SPPB rutin", "Alur rutin dasar", "Tab aktif/riwayat, satu siklus mingguan, dan barang tambahan setelah SPPB selesai"),
            ("SPPB non rutin", "Persetujuan berjenjang awal", "Status operasional lebih lengkap sampai dana, logistik, dan siap diserahkan"),
            ("Barang tambahan", "Belum ada pengelolaan khusus", "Tambah/edit/hapus item pada SPPB lama dengan label Tambahan di seluruh role"),
            ("Stok", "Kartu stok dan mutasi dasar", "Mutasi desc 20 data, filter semua barang, PDF portrait, dan pencatatan stok SPPB"),
            ("Master inventaris", "Kategori/KIB aset bawaan", "Hierarki Unit/Kategori/Kelompok/Jenis/Barang dan nomor inventaris otomatis"),
            ("Registrasi aset", "KIB A-F aktif", "Menggunakan Jenis Inventaris terbaru; KIB lama hanya kompatibilitas arsip"),
            ("Hak akses", "Role Kanit/Kasie/Kabid dasar", "Permission per endpoint dan pengamanan manajemen modul"),
            ("Notifikasi", "Belum ada arsitektur lengkap", "Database + polling + Workerman + Pusher opsional"),
            ("Frontend", "Handler JS modul", "Overlay proses, anti-submit ganda, JSON error handling, breadcrumb, dan layout responsif"),
            ("Laporan", "Laporan dasar", "Kartu stok PDF, cost unit, pengadaan, distribusi, aset, dan audit lebih luas"),
        ],
        [3.1, 5.7, 7.9],
    )

    add_title(doc, "2. Method Baru yang Menjadi Pembeda")
    doc.add_paragraph("Admin.php RSNS menambahkan 52 method. Kelompok terpenting:")
    add_table(
        doc,
        ["Kelompok", "Method utama", "Fungsi"],
        [
            ("SPPB tambahan", "postTambah/Update/HapusItemSppbTambahan, _canManageSppbTambahan", "Menempelkan item tambahan pada nomor SPPB rutin yang sudah ada"),
            ("Alur SPPB", "_renderSppbPage, postProsesLogistikSppb, postKeputusanDanaSppb", "Memisahkan rutin/non rutin dan tahapan sesuai role"),
            ("Stok SPPB", "_catatStokKeluarSppb, _getSppbStokKurang, _updateStokKeluarTanpaLokasi", "Validasi dan pencatatan barang keluar satu kali"),
            ("Inventaris", "get/post MasterInventaris, _generateNomorInventaris", "Master terbaru dan nomor inventaris berjenjang"),
            ("Mutasi/Kartu stok", "postImportMutasiCsv, getExportMutasiCsv, anyCetakKartuStokPdf", "Impor, ekspor, dan cetak kartu stok"),
            ("Notifikasi", "_broadcastNotification, _triggerPusherNotification, anyCekNotifikasi", "Database, queue Workerman, Pusher, dan polling"),
            ("Hak akses", "_getAllChildUnitKodes, _syncLogistikAccess", "Pembatasan unit dan sinkronisasi akses modul"),
            ("Perencanaan", "getRencanapembelian, getPermintaanNonRutin, getRealisasibelanja", "Perencanaan dan realisasi pengadaan baru"),
        ],
        [3.2, 6.5, 7.0],
    )

    add_title(doc, "3. Klasifikasi File Update")
    add_table(
        doc,
        ["Kelas", "File/Folder", "Tindakan"],
        [
            ("Wajib", "plugins/logistik_non_medis/Admin.php", "Merge manual; pertahankan seluruh endpoint, migrasi, status, stok, aset, dan notifikasi RSNS"),
            ("Wajib", "plugins/logistik_non_medis/js/admin/logistik.js", "Pakai versi RSNS agar overlay proses, anti-double submit, dan parsing JSON tetap bekerja"),
            ("Wajib", "plugins/logistik_non_medis/css/admin/logistik.css", "Pakai bersama view terbaru agar form/modal tidak terpotong"),
            ("Wajib", "plugins/logistik_non_medis/view/admin/", "Deploy sebagai satu pasangan dengan Admin.php; endpoint dan selector saling bergantung"),
            ("Wajib", "workerman.php", "Dibutuhkan bila mode realtime Workerman digunakan"),
            ("Wajib", "composer.json + composer.lock/vendor hasil install", "Workerman dipatok 4.0.47 dan platform PHP 7.4.33"),
            ("Wajib", "service-worker.js", "Mencegah halaman admin dinavigasikan melalui cache dan menghapus cache lama"),
            ("Tinjau", "systems/lib/QueryWrapper.php", "Menghapus ONLY_FULL_GROUP_BY pada sesi DB; berdampak ke seluruh aplikasi"),
            ("Tinjau", "plugins/modules/Admin.php", "Pengamanan menu install/uninstall modul hanya untuk administrator"),
            ("Tinjau", "plugins/permintaan_logistik_non_medis/", "Modul tambahan; aktifkan hanya bila masih dipakai, agar tidak menggandakan menu SPPB"),
            ("Opsional", "aset-info.php dan aset-qr.php", "Endpoint root untuk QR; audit autentikasi dan data yang ditampilkan sebelum publik"),
            ("SQL", "schema_procurement.sql, hakakses_migration.sql", "Jalankan sesuai kebutuhan setelah backup dan pemeriksaan struktur"),
        ],
        [2.2, 6.3, 8.2],
    )

    doc.add_paragraph("File yang tidak boleh masuk paket produksi:")
    add_table(
        doc,
        ["File", "Risiko"],
        [
            ("admin/debug_sppb.txt", "Data debug tidak diperlukan dan dapat membuka detail internal"),
            ("check_pass.php", "Alat pemeriksaan password tidak layak berada pada web root produksi"),
            ("plugins/logistik_non_medis/truncate_unit.php", "Berpotensi mengosongkan data unit"),
            ("Data_Aset_Non_Medis_*.csv dan agent/*.csv/*.xlsx", "Data operasional/contoh tidak boleh otomatis disalin"),
            ("migrate_hapus_kolom_kategori_aset.sql", "Menghapus kolom; tidak sesuai prinsip mempertahankan data"),
            ("assign_pj.sql", "Data assignment spesifik; hanya dijalankan bila memang ditujukan ke server tersebut"),
            ("dump database dan folder agent/db_backups", "Backup tidak boleh menjadi bagian source release"),
        ],
        [7.6, 9.1],
    )
    add_callout(doc, "Info.php perlu dibersihkan", "File Info.php saat ini memiliki key category dua kali. PHP memakai nilai terakhir, tetapi duplikasi sebaiknya dihapus sebelum rilis agar metadata modul tidak membingungkan.", "warning")

    add_title(doc, "4. Peta View yang Harus Sinkron")
    add_table(
        doc,
        ["Kelompok view", "File penting", "Ketergantungan"],
        [
            ("Dashboard", "manage.html, notifikasi.html", "Menu role, kartu ringkasan, polling/realtime"),
            ("SPPB", "distribusi.sppb.html/form/display/detail/cetak", "Status enum, barang tambahan, endpoint approval, pengambilan"),
            ("Gudang", "gudang.stok.*, gudang.mutasi.*, gudang.opnamev2.*", "Endpoint data, pagination, CSV, PDF"),
            ("Master", "master.barang.*, master.inventaris.*", "Struktur kategori/kelompok/jenis/barang"),
            ("Aset", "aset.registrasi.*, aset.kib.*, aset.mutasi.*", "Master inventaris, QR, nomor inventaris, breadcrumb"),
            ("Pengadaan", "pengadaan.rencana_*, terima_rutin.*, realisasi.html", "Enam tabel schema_procurement.sql"),
            ("Laporan", "laporan.costunit.*, distribusi, pengadaan, aset", "Query agregasi dan mode SQL database"),
        ],
        [3.0, 7.5, 6.2],
    )
    add_callout(doc, "Aturan merge view", "Jangan memasang satu view baru tanpa Admin.php dan JavaScript pasangannya. Perbedaan nama endpoint/ID selector akan menghasilkan modal kosong, redirect salah, atau respons HTML yang gagal diparse sebagai JSON.", "danger")

    add_title(doc, "5. Perbedaan Struktur Database")
    add_callout(doc, "Sumber pemeriksaan", "Struktur berikut dibandingkan dari Admin.php dan database lokal mlite_rsns pada 02 Agustus 2026. Struktur server wajib diperiksa ulang melalui information_schema sebelum migrasi.", "info")
    add_table(
        doc,
        ["Tabel baru terhadap Rifangga", "Kolom inti", "Pemicu"],
        [
            ("..._inventaris_master", "jenis_master, kode, kode_inventaris, kategori, kelompok, jenis, barang, nama", "_initAset / Master Inventaris"),
            ("..._inventaris_kategori", "kode_kategori, nama_kategori", "_initAset"),
            ("..._inventaris_kelompok", "kategori + kelompok + nama", "_initAset"),
            ("..._inventaris_jenis", "kategori + kelompok + jenis + nama", "_initAset"),
            ("..._notifikasi", "target, pesan, tipe, url, is_read, tanggal", "_initNotifikasi"),
        ],
        [6.7, 6.5, 3.5],
    )

    doc.add_paragraph("Tabel yang berubah paling besar:")
    add_table(
        doc,
        ["Tabel", "Penyesuaian RSNS"],
        [
            ("..._sppb", "47 kolom aktual; jenis_permintaan, jenis_keluar, item manual, data tujuan, approval, cost, verifikasi, pengambil, dan status panjang"),
            ("..._aset", "60 kolom aktual; nomor inventaris, dokumen, lokasi fisik, bahan, tahun, satuan, jumlah, penyusutan, dan kolom KIB legacy"),
            ("..._user_roles", "kode_unit pada database lokal bertipe text agar satu user dapat menangani lebih dari satu unit"),
            ("..._master_barang", "kategori/kode kategori, tipe barang, jenis item, stok minimum/maksimum/safety, dan lokasi default"),
            ("..._kartu_stok", "batch, harga, dan tipe transaksi yang diperluas"),
            ("mlite_settings", "Empat setting realtime baru disimpan per module/field"),
        ],
        [5.0, 11.7],
    )

    add_title(doc, "6. SQL Struktur Baru")
    add_code(
        doc,
        """
CREATE TABLE IF NOT EXISTS rsns_custom_logistik_non_medis_notifikasi (
  id INT NOT NULL AUTO_INCREMENT,
  user_target VARCHAR(100) NOT NULL,
  pesan TEXT NOT NULL,
  tipe VARCHAR(50) DEFAULT NULL,
  url VARCHAR(255) DEFAULT NULL,
  is_read TINYINT(1) NOT NULL DEFAULT 0,
  tgl_dibuat DATETIME NOT NULL,
  PRIMARY KEY (id),
  KEY user_target (user_target),
  KEY is_read (is_read)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
""",
        "Tabel notifikasi",
    )
    add_code(
        doc,
        """
CREATE TABLE IF NOT EXISTS rsns_custom_logistik_non_medis_inventaris_kategori (
  kode_kategori CHAR(1) NOT NULL,
  nama_kategori VARCHAR(100) NOT NULL,
  PRIMARY KEY (kode_kategori)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;

CREATE TABLE IF NOT EXISTS rsns_custom_logistik_non_medis_inventaris_kelompok (
  kode_kategori CHAR(1) NOT NULL,
  kode_kelompok CHAR(2) NOT NULL,
  nama_kelompok VARCHAR(150) NOT NULL,
  PRIMARY KEY (kode_kategori, kode_kelompok)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;

CREATE TABLE IF NOT EXISTS rsns_custom_logistik_non_medis_inventaris_jenis (
  kode_kategori CHAR(1) NOT NULL,
  kode_kelompok CHAR(2) NOT NULL,
  kode_jenis CHAR(2) NOT NULL,
  nama_jenis VARCHAR(150) NOT NULL,
  PRIMARY KEY (kode_kategori, kode_kelompok, kode_jenis)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
""",
        "Hierarki inventaris terbaru",
    )
    add_code(
        doc,
        """
CREATE TABLE IF NOT EXISTS rsns_custom_logistik_non_medis_inventaris_master (
  id INT NOT NULL AUTO_INCREMENT,
  jenis_master ENUM('UNIT','BARANG') NOT NULL,
  kode VARCHAR(50) NOT NULL,
  kode_inventaris CHAR(3) DEFAULT NULL,
  kode_kategori CHAR(1) NOT NULL DEFAULT '',
  nama VARCHAR(200) NOT NULL,
  kode_kelompok CHAR(2) DEFAULT NULL,
  kode_jenis CHAR(2) DEFAULT NULL,
  kode_barang CHAR(2) DEFAULT NULL,
  nama_kelompok VARCHAR(150) DEFAULT NULL,
  nama_jenis VARCHAR(150) DEFAULT NULL,
  kib_jenis ENUM('A','B','C','D','E','F') DEFAULT NULL,
  status ENUM('Aktif','Nonaktif') NOT NULL DEFAULT 'Aktif',
  tgl_input DATETIME DEFAULT NULL,
  PRIMARY KEY (id),
  UNIQUE KEY jenis_kode (jenis_master, kode_kategori, kode)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
""",
        "Master inventaris",
    )
    add_callout(doc, "KIB lama", "Kolom kib_jenis tetap ada untuk kompatibilitas data lama, tetapi form registrasi terbaru memakai Kelompok > Jenis > Barang. Jangan menjalankan migrasi yang menghapus data KIB tanpa audit.", "warning")

    add_title(doc, "7. Kolom SPPB yang Wajib Tersedia")
    add_table(
        doc,
        ["Kelompok", "Kolom"],
        [
            ("Jenis", "minggu_ke, jenis_permintaan, jenis_keluar"),
            ("Item", "item_sumber, nama_barang_manual, spesifikasi_manual, estimasi_harga"),
            ("Tujuan", "latar_belakang_tujuan, sasaran_kegunaan, rencana_digunakan, sifat_permintaan"),
            ("Penanggung jawab", "diajukan_oleh, penanggung_jawab_1, penanggung_jawab_2, ka_unit"),
            ("Approval", "user/tgl approve Ka Unit, Ka Sie, Ka Bidang, dan Unit"),
            ("Logistik", "jumlah_disetujui, user/tgl verifikasi, keterangan_verifikasi"),
            ("Cost", "harga_satuan_cost, subtotal_cost, user_cost, tgl_cost"),
            ("Penyelesaian", "diambil_oleh, tgl_diambil, ditolak_pada_status"),
        ],
        [4.7, 12.0],
    )
    add_code(
        doc,
        """
SELECT COLUMN_NAME, COLUMN_TYPE, IS_NULLABLE, COLUMN_DEFAULT
FROM information_schema.COLUMNS
WHERE TABLE_SCHEMA = DATABASE()
  AND TABLE_NAME = 'rsns_custom_logistik_non_medis_sppb'
ORDER BY ORDINAL_POSITION;

SELECT status, jenis_permintaan, jenis_keluar, COUNT(*) jumlah
FROM rsns_custom_logistik_non_medis_sppb
GROUP BY status, jenis_permintaan, jenis_keluar;
""",
        "Verifikasi sebelum mengubah enum/status",
    )
    add_callout(doc, "Jangan ALTER enum tanpa audit", "Nilai enum yang tidak tercantum dapat dipotong menjadi string kosong. Simpan hasil GROUP BY, backup tabel, lalu pastikan enum baru mencakup semua status lama dan baru.", "danger")

    add_title(doc, "8. SQL Pengadaan Tambahan")
    doc.add_paragraph("File schema_procurement.sql menambahkan enam tabel:")
    add_bullet(doc, "rencana_rutin dan rencana_rutin_detail")
    add_bullet(doc, "terima_rutin dan terima_rutin_detail")
    add_bullet(doc, "rencana_nonrutin dan rencana_nonrutin_detail")
    add_callout(doc, "Schema file tertinggal dari database lokal", "Database lokal sudah memiliki status Ditolak, alasan_penolakan, dan kode_unit pada beberapa tabel, sedangkan schema_procurement.sql belum memuat semuanya. Jangan memakai file itu sebagai satu-satunya sumber migrasi server; samakan dengan struktur lokal yang sudah diuji.", "warning")
    add_code(
        doc,
        """
SHOW CREATE TABLE rsns_custom_logistik_non_medis_rencana_rutin;
SHOW CREATE TABLE rsns_custom_logistik_non_medis_rencana_nonrutin;
SHOW CREATE TABLE rsns_custom_logistik_non_medis_terima_rutin;
""",
        "Ambil struktur acuan dari database staging yang sudah benar",
    )

    add_title(doc, "9. Urutan Migrasi Database yang Disarankan")
    add_number(doc, "Backup database penuh dan catat jumlah baris tabel SPPB, kartu stok, mutasi, aset, serta master inventaris.")
    add_number(doc, "Pasang kode pada staging dan jalankan composer install.")
    add_number(doc, "Jalankan hakakses_migration.sql; skrip ini menggunakan CREATE IF NOT EXISTS dan INSERT IGNORE.")
    add_number(doc, "Bandingkan schema_procurement.sql dengan SHOW CREATE TABLE dari database lokal sebelum menjalankannya.")
    add_number(doc, "Buka menu SPPB, Gudang, Master Inventaris, Registrasi Aset, dan Notifikasi agar fungsi _init... berjalan.")
    add_number(doc, "Periksa log SQL/PHP dan ulangi query information_schema.")
    add_number(doc, "Tambahkan indeks hanya jika belum ada; performance_indexes.sql tidak idempotent dan akan error bila nama indeks sudah dipakai.")
    add_number(doc, "Bandingkan kembali jumlah baris. Migrasi struktur tidak boleh mengurangi data.")
    add_callout(doc, "Migrasi aset berisiko", "_initAset saat ini memiliki DROP TABLE IF EXISTS aset_mutasi ketika kolom no_mutasi tidak ditemukan. Sebelum deploy, ubah menjadi migrasi ALTER/penyalinan data yang aman. Jangan memicu menu aset pada database produksi lama sebelum bagian ini diperbaiki atau tabel sudah dibackup dan dimigrasikan.", "danger")

    add_title(doc, "10. Workerman: Perbedaan dan Penyesuaian")
    doc.add_paragraph("Alur saat ini:")
    add_code(
        doc,
        """
Admin.php menyimpan notifikasi ke database
  -> menulis JSON per baris ke tmp/logistik_notifications.queue
  -> workerman.php membaca queue setiap 0,5 detik
  -> broadcast WebSocket port 3892
  -> browser menerima event lalu mengambil data aman dari endpoint ceknotifikasi
  -> polling 10 detik tetap menjadi fallback
""",
    )
    add_table(
        doc,
        ["File", "Penyesuaian terhadap Rifangga"],
        [
            ("composer.json", "workerman/workerman dipatok 4.0.47; platform PHP dipatok 7.4.33"),
            ("workerman.php", "Tambah Timer, queue file, offset, dan broadcast event logistik_notification"),
            ("Admin.php", "_broadcastNotification menulis queue dengan FILE_APPEND + LOCK_EX"),
            ("notifikasi.html", "WebSocket ws/wss port 3892, maksimal dua retry, lalu polling"),
        ],
        [5.0, 11.7],
    )
    add_code(
        doc,
        """
composer install --no-dev --prefer-dist --optimize-autoloader
mkdir -p tmp
chown -R www-data:www-data tmp
chmod -R 775 tmp
php workerman.php start -d
php workerman.php status
""",
        "Aktivasi Workerman di server Linux",
    )
    add_callout(doc, "Dua risiko queue saat ini", "Offset dimulai dari 0 sehingga event lama dapat diputar ulang setelah worker restart. Queue juga terus membesar. Disarankan mulai dari filesize(queue) saat startup dan lakukan rotasi terkontrol ketika worker berhenti.", "warning")
    add_callout(doc, "Privasi broadcast", "Worker mengirim payload target/pesan ke semua koneksi. Untuk server internet, broadcast hanya sinyal type/id atau tambahkan autentikasi dan pemisahan koneksi per target. Browser tetap mengambil isi notifikasi melalui endpoint yang difilter role/unit.", "danger")

    add_title(doc, "11. HTTP, HTTPS, dan WSS")
    add_table(
        doc,
        ["Lingkungan", "Pilihan"],
        [
            ("Intranet HTTP", "Workerman langsung ws://HOST:3892; buka firewall hanya untuk jaringan internal"),
            ("Internet HTTPS", "Pilih Pusher atau terminasi TLS/reverse proxy WSS; browser tidak boleh memakai ws pada halaman HTTPS"),
            ("Shared hosting", "Gunakan Pusher + polling karena proses daemon/port biasanya tidak tersedia"),
            ("VPS HTTPS", "Workerman bind localhost, Nginx/Apache meneruskan wss://domain:3892 ke port internal"),
        ],
        [4.2, 12.5],
    )
    add_code(
        doc,
        """
# Penyesuaian workerman.php untuk reverse proxy:
$ws = new Worker('websocket://127.0.0.1:3893');

# Nginx public listener 3892 memakai sertifikat yang sama dengan aplikasi,
# kemudian proxy_pass ke http://127.0.0.1:3893 dan meneruskan Upgrade/Connection.
# notifikasi.html dapat tetap memakai wss://domain:3892.
""",
        "Model kompatibel dengan URL JavaScript saat ini",
    )

    add_title(doc, "12. Pusher: Struktur dan Konfigurasi")
    doc.add_paragraph("Pusher aktif hanya bila setting notif_realtime bernilai pusher dan seluruh credential terisi. Backend membuat signature REST sendiri dengan cURL; tidak memakai SDK PHP Pusher.")
    add_table(
        doc,
        ["mlite_settings.field", "Nilai", "Dikirim ke browser"],
        [
            ("notif_realtime", "workerman atau pusher", "Ya, mode saja"),
            ("pusher_app_id", "App ID", "Tidak"),
            ("pusher_key", "Public key", "Ya"),
            ("pusher_secret", "Secret", "Tidak"),
            ("pusher_cluster", "Contoh ap1", "Ya"),
        ],
        [5.4, 5.4, 5.9],
    )
    add_code(
        doc,
        """
INSERT INTO mlite_settings (module, field, value)
SELECT 'logistik_non_medis', 'notif_realtime', 'pusher'
WHERE NOT EXISTS (
  SELECT 1 FROM mlite_settings
  WHERE module='logistik_non_medis' AND field='notif_realtime'
);

UPDATE mlite_settings SET value='PUSHER_APP_ID'
WHERE module='logistik_non_medis' AND field='pusher_app_id';
UPDATE mlite_settings SET value='PUSHER_PUBLIC_KEY'
WHERE module='logistik_non_medis' AND field='pusher_key';
UPDATE mlite_settings SET value='PUSHER_SECRET'
WHERE module='logistik_non_medis' AND field='pusher_secret';
UPDATE mlite_settings SET value='ap1'
WHERE module='logistik_non_medis' AND field='pusher_cluster';
""",
        "Contoh konfigurasi - ganti placeholder dan jangan commit credential",
    )
    add_callout(doc, "Pastikan record setting tersedia", "Buka endpoint notifikasi sekali agar _initNotifikasi membuat empat record default, atau gunakan INSERT NOT EXISTS untuk setiap field sebelum UPDATE.", "info")
    add_callout(doc, "Channel Pusher saat ini publik", "Nama channel berbentuk logistik-<target>. Payload memuat pesan. Untuk data sensitif gunakan private channel dengan endpoint autentikasi, atau ubah event menjadi sinyal tanpa isi pesan lalu biarkan browser mengambil isi melalui ceknotifikasi.", "danger")

    add_title(doc, "13. Penyesuaian File yang Direkomendasikan")
    add_table(
        doc,
        ["Prioritas", "File", "Penyesuaian sebelum produksi"],
        [
            ("P0", "Admin.php::_initAset", "Hilangkan DROP TABLE aset_mutasi; migrasikan kolom secara additive dan salin data lama"),
            ("P0", "workerman.php", "Jangan replay queue lama; mulai offset di EOF atau simpan offset persisten"),
            ("P0", "workerman.php/Admin.php", "Jangan broadcast isi pesan ke semua klien; kirim event minimal"),
            ("P0", "Release manifest", "Keluarkan check_pass.php, debug file, truncate_unit.php, CSV/XLSX, dan dump DB"),
            ("P1", "notifikasi.html", "Buat host/port/path WebSocket configurable; dukung reverse proxy HTTPS"),
            ("P1", "schema_procurement.sql", "Samakan enum Ditolak, alasan_penolakan, dan kode_unit dengan database terbaru"),
            ("P1", "performance_indexes.sql", "Tambahkan pemeriksaan information_schema agar idempotent"),
            ("P1", "Info.php", "Hapus key category yang duplikat"),
            ("P2", "QueryWrapper.php", "Pertimbangkan perbaikan query GROUP BY daripada mematikan ONLY_FULL_GROUP_BY secara global"),
            ("P2", "service-worker.js", "Naikkan cacheName jika daftar aset statis berubah"),
        ],
        [2.0, 5.1, 9.6],
    )

    add_title(doc, "14. Urutan Penerapan Paket Update")
    add_number(doc, "Rapi dan commit semua perubahan RSNS yang memang akan dipakai.")
    add_number(doc, "Buat manifest hanya file produksi berdasarkan Bab 3.")
    add_number(doc, "Lakukan penyesuaian P0 sebelum paket dianggap siap.")
    add_number(doc, "Bandingkan skema server dengan database staging memakai information_schema/SHOW CREATE TABLE.")
    add_number(doc, "Backup server, pasang kode, lalu composer install.")
    add_number(doc, "Migrasikan hak akses, SPPB, inventaris, pengadaan, aset, dan notifikasi secara bertahap.")
    add_number(doc, "Pilih satu mode realtime utama: Workerman atau Pusher. Polling tetap aktif sebagai fallback.")
    add_number(doc, "Uji role unit, Kanit, Kasie, Kabid, admin, gudang, logistik, dan aset.")
    add_number(doc, "Verifikasi jumlah data sebelum/sesudah dan simpan hasil sebagai bukti update.")

    add_title(doc, "15. Checklist Verifikasi Delta")
    add_checklist(
        doc,
        [
            "rifangga/main 027fcfd sudah ancestor paket RSNS.",
            "Paket tidak memuat file debug, credential, dump, truncate, CSV, atau XLSX operasional.",
            "Admin.php, view, JS, dan CSS berasal dari satu versi yang sama.",
            "Lima tabel runtime baru tersedia.",
            "SPPB mempunyai semua kolom jenis, tambahan, approval, cost, dan pengambilan.",
            "Enum status server mencakup seluruh nilai yang sudah tersimpan.",
            "Master Inventaris Kelompok/Jenis/Barang tersedia dan aset lama tidak dihapus.",
            "Migrasi aset tidak menjalankan DROP TABLE pada data lama.",
            "tmp dapat ditulis web server dan worker.",
            "Workerman/Pusher hanya mengirim data sesuai kebijakan keamanan.",
            "Polling 10 detik tetap bekerja saat realtime dimatikan.",
            "Jumlah SPPB, kartu stok, mutasi, aset, dan master sebelum/sesudah sama atau bertambah wajar.",
        ],
    )

    add_title(doc, "16. Perintah Pembanding untuk Update Berikutnya")
    add_code(
        doc,
        """
git fetch rifangga main --prune
git log --oneline HEAD..rifangga/main
git diff --stat rifangga/main
git diff --name-status rifangga/main
git diff rifangga/main -- plugins/logistik_non_medis/Admin.php
git diff rifangga/main -- workerman.php service-worker.js composer.json
git merge-base --is-ancestor rifangga/main HEAD
""",
    )
    add_callout(
        doc,
        "Kesimpulan",
        "Sistem RSNS bukan sekadar Rifangga terbaru yang disalin ke server. Ia adalah turunan dengan 52 method baru, struktur inventaris dan notifikasi baru, alur SPPB yang diperluas, serta konfigurasi realtime. Update paling aman dilakukan sebagai paket delta terverifikasi dengan migrasi additive dan manifest file produksi.",
        "success",
    )

    doc.core_properties.title = "Delta Update Rifangga ke Sistem RSNS Terbaru"
    doc.core_properties.subject = "Perbedaan file, SQL, Workerman, Pusher, dan penyesuaian update"
    doc.core_properties.author = "Tim Sistem Informasi RSNS"
    doc.core_properties.modified = datetime.now()
    doc.save(OUTPUT_FILE)
    return OUTPUT_FILE


if __name__ == "__main__":
    print(build_document())
