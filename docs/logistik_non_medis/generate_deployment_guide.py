from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT_DIR = Path(__file__).resolve().parent
OUTPUT_FILE = OUTPUT_DIR / "04_Modul_Deploy_dan_Update_Aman_Rifangga.docx"

DOC_DATE = "02 Agustus 2026"
NAVY = "34495E"
BLUE = "3C78A8"
GREEN = "5B9B62"
ORANGE = "E59B35"
RED = "C74B50"
TEXT = "3E4852"
WHITE = "FFFFFF"
LIGHT_BLUE = "EAF3F8"
LIGHT_GREEN = "EDF6EE"
LIGHT_ORANGE = "FFF4E1"
LIGHT_RED = "FCEBEC"
LIGHT_GRAY = "F2F4F5"
MID_GRAY = "D5DADF"
CODE_BG = "202830"


def shade(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:fill"), color)


def margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        item = tc_mar.find(qn(f"w:{name}"))
        if item is None:
            item = OxmlElement(f"w:{name}")
            tc_mar.append(item)
        item.set(qn("w:w"), str(value))
        item.set(qn("w:type"), "dxa")


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Halaman ")
    run.font.size = Pt(8)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def configure(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.75)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in (
        ("Title", 24, NAVY),
        ("Heading 1", 16, NAVY),
        ("Heading 2", 12.5, BLUE),
        ("Heading 3", 10.5, TEXT),
    ):
        style = doc.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    header = section.header.paragraphs[0]
    header.text = "PANDUAN OPERASIONAL | DEPLOY DAN UPDATE AMAN MLITE RSNS"
    header.runs[0].font.name = "Arial"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.bold = True
    header.runs[0].font.color.rgb = RGBColor.from_string(NAVY)
    p_pr = header._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:color"), BLUE)
    border.append(bottom)
    p_pr.append(border)

    page_number(section.footer.paragraphs[0])
    props = doc.core_properties
    props.title = "Modul Deploy dan Update Aman mLITE RSNS"
    props.subject = "Integrasi commit Rifangga, release, deployment server, dan rollback"
    props.author = "Tim Sistem Informasi RSNS"
    props.keywords = "mLITE, logistik non medis, git, deploy, Rifangga, rollback"


def add_title(doc, title, subtitle=None):
    p = doc.add_paragraph(title, style="Heading 1")
    if subtitle:
        q = doc.add_paragraph(subtitle)
        q.runs[0].font.italic = True
        q.runs[0].font.color.rgb = RGBColor.from_string(BLUE)
    return p


def add_bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    if bold_prefix and text.startswith(bold_prefix):
        p.add_run(bold_prefix).bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def add_callout(doc, title, text, kind="info"):
    palette = {
        "info": (LIGHT_BLUE, BLUE),
        "success": (LIGHT_GREEN, GREEN),
        "warning": (LIGHT_ORANGE, ORANGE),
        "danger": (LIGHT_RED, RED),
    }
    background, accent = palette[kind]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(0.28)
    table.columns[1].width = Cm(16.6)
    shade(table.cell(0, 0), accent)
    shade(table.cell(0, 1), background)
    margins(table.cell(0, 0), 80, 20, 80, 20)
    margins(table.cell(0, 1), 130, 180, 130, 180)
    p = table.cell(0, 1).paragraphs[0]
    r = p.add_run(title + "\n")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(accent)
    p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_code(doc, code, title=None):
    if title:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(9)
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Cm(16.8)
    cell = table.cell(0, 0)
    shade(cell, CODE_BG)
    margins(cell, 130, 160, 130, 160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for index, line in enumerate(code.strip().splitlines()):
        if index:
            p.add_run().add_break()
        run = p.add_run(line.rstrip())
        run.font.name = "Consolas"
        run.font.size = Pt(8.2)
        run.font.color.rgb = RGBColor.from_string("E8EEF2")
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    for i, value in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, NAVY)
        margins(cell)
        p = cell.paragraphs[0]
        r = p.add_run(str(value))
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(WHITE)
        if widths:
            cell.width = Cm(widths[i])
    repeat_header(table.rows[0])
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(values):
            if row_index % 2:
                shade(cells[i], LIGHT_GRAY)
            margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells[i].text = str(value)
            if widths:
                cells[i].width = Cm(widths[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_checklist(doc, items):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.35)
        p.paragraph_format.first_line_indent = Cm(-0.35)
        p.add_run("[ ] ").bold = True
        p.add_run(item)


def add_flow(doc):
    labels = [
        ("RIFANGGA", "Sumber update", LIGHT_ORANGE, ORANGE),
        ("BRANCH INTEGRASI", "Gabung + konflik", LIGHT_BLUE, BLUE),
        ("STAGING", "Uji alur + DB", LIGHT_GREEN, GREEN),
        ("ORIGIN + TAG", "Rilis milik RSNS", LIGHT_BLUE, BLUE),
        ("SERVER", "Deploy + smoke test", LIGHT_GREEN, GREEN),
    ]
    table = doc.add_table(rows=1, cols=9)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, (title, subtitle, fill, accent) in enumerate(labels):
        cell = table.cell(0, index * 2)
        shade(cell, fill)
        margins(cell, 130, 60, 130, 60)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title + "\n")
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor.from_string(accent)
        s = p.add_run(subtitle)
        s.font.size = Pt(7.5)
        if index < len(labels) - 1:
            arrow = table.cell(0, index * 2 + 1)
            arrow.text = ">"
            arrow.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            arrow.paragraphs[0].runs[0].font.bold = True
            arrow.paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(NAVY)
    doc.add_paragraph()


def build_document():
    doc = Document()
    configure(doc)

    p = doc.add_paragraph("DOKUMEN INTERNAL", style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\n")
    title = doc.add_paragraph("MODUL DEPLOY DAN UPDATE AMAN", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    product = doc.add_paragraph("mLITE RSNS - LOGISTIK NON MEDIS")
    product.alignment = WD_ALIGN_PARAGRAPH.CENTER
    product.runs[0].font.name = "Arial"
    product.runs[0].font.size = Pt(18)
    product.runs[0].font.bold = True
    product.runs[0].font.color.rgb = RGBColor.from_string(BLUE)

    band = doc.add_table(rows=1, cols=1)
    band.alignment = WD_TABLE_ALIGNMENT.CENTER
    band.autofit = False
    band.columns[0].width = Cm(15.6)
    cell = band.cell(0, 0)
    shade(cell, NAVY)
    margins(cell, 260, 250, 260, 250)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Integrasi commit Rifangga, penerbitan rilis, deployment server, dan rollback")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor.from_string(WHITE)

    doc.add_paragraph("\n")
    info = add_table(
        doc,
        ["Identitas", "Keterangan"],
        [
            ("Versi dokumen", "1.0"),
            ("Tanggal snapshot", DOC_DATE),
            ("Sasaran", "Rifangga / pengelola aplikasi / administrator server"),
            ("Repo rilis RSNS", "https://github.com/Hanaf1/LOGISTIKRSNSMLITE.git"),
            ("Repo sumber update", "https://github.com/Rifangga99/mlite_rsns.git"),
        ],
        [4.3, 12.4],
    )
    doc.add_paragraph("Dokumen ini dibuat untuk menjaga penyesuaian RSNS yang sudah berjalan agar tidak hilang ketika mengambil commit baru atau melakukan deploy ke server.").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    add_title(doc, "Ringkasan Eksekutif")
    add_callout(
        doc,
        "Status saat dokumen dibuat",
        "Commit terbaru rifangga/main adalah 027fcfd (24 Juli 2026) dan sudah menjadi ancestor branch lokal. Branch lokal berada 23 commit di depan rifangga/main dan 25 commit di depan origin/main. Tidak ada commit Rifangga baru yang perlu digabung saat snapshot ini dibuat.",
        "success",
    )
    add_callout(
        doc,
        "Repo lokal belum siap langsung dideploy",
        "Terdapat 27 file termodifikasi dan 10 item belum terlacak. Semua perubahan yang memang diperlukan harus diperiksa lalu dikomit terlebih dahulu. Jangan memakai git add . karena ada backup database dan dokumen kerja yang tidak semestinya ikut ke rilis.",
        "warning",
    )
    doc.add_paragraph("Aturan inti deployment:")
    add_bullet(doc, "Rifangga adalah sumber pembaruan, bukan sumber deploy server.")
    add_bullet(doc, "Server mengambil kode dari origin milik RSNS pada commit/tag yang sudah diuji.")
    add_bullet(doc, "Database, config.php, uploads, dan data operasional tidak boleh ditimpa oleh update kode.")
    add_bullet(doc, "Setiap deploy wajib punya backup database, backup file penting, dan catatan commit sebelum deploy.")
    add_bullet(doc, "Konflik pada Admin.php, JavaScript, CSS, dan view diselesaikan per fungsi; jangan mengambil seluruh versi ours/theirs.")
    add_flow(doc)

    add_title(doc, "1. Tujuan dan Ruang Lingkup")
    doc.add_paragraph("Modul ini dipakai untuk empat skenario:")
    add_number(doc, "Merapikan perubahan lokal menjadi commit rilis yang dapat dilacak.")
    add_number(doc, "Mengambil commit baru dari Rifangga tanpa mengurangi fitur dan data RSNS.")
    add_number(doc, "Menerbitkan commit yang sudah diuji ke origin dan memberi tag rilis.")
    add_number(doc, "Memasang rilis ke server serta mengembalikan versi lama bila terjadi masalah.")
    add_callout(doc, "Bukan panduan reset", "Dokumen ini tidak pernah mengandalkan git reset --hard, force push, atau penggantian database produksi. Operasi tersebut berisiko menghilangkan pekerjaan dan data.", "danger")

    add_title(doc, "2. Peta Repository dan Branch")
    add_table(
        doc,
        ["Nama", "Nilai saat snapshot", "Fungsi"],
        [
            ("Remote origin", "Hanaf1/LOGISTIKRSNSMLITE", "Sumber rilis yang dipakai server"),
            ("Remote rifangga", "Rifangga99/mlite_rsns", "Sumber pembaruan upstream"),
            ("Branch kerja", "fitur-hak-akses-style-lama", "Tempat penyesuaian RSNS saat ini"),
            ("HEAD lokal", "39a5024", "Merge Rifangga terakhir ke branch kerja"),
            ("origin/main", "845dc7c", "Main RSNS sebelum perubahan branch kerja diterbitkan"),
            ("rifangga/main", "027fcfd", "Commit upstream terbaru saat snapshot"),
        ],
        [3.2, 5.8, 7.7],
    )
    add_code(
        doc,
        """
git remote -v
git branch -vv
git fetch --all --prune
git log -5 --oneline --decorate
git rev-list --left-right --count HEAD...rifangga/main
git merge-base --is-ancestor rifangga/main HEAD
""",
        "Periksa ulang kondisi terbaru",
    )
    add_callout(doc, "Membaca hasil ancestor", "Jika perintah terakhir selesai tanpa pesan dan exit code 0, commit rifangga/main sudah terkandung di branch lokal. Bila tidak, lakukan proses integrasi pada Bab 5.", "info")

    add_title(doc, "3. Data yang Wajib Dilindungi")
    add_table(
        doc,
        ["Komponen", "Perlakuan", "Alasan"],
        [
            ("Database MySQL/MariaDB", "Backup penuh sebelum deploy", "Memuat transaksi SPPB, stok, mutasi, aset, user, dan konfigurasi"),
            ("config.php", "Pertahankan versi server", "Berisi koneksi DB dan pengaturan lingkungan; file di-ignore Git"),
            ("uploads/", "Jangan hapus atau replace", "Berisi foto, dokumen, bukti, dan lampiran"),
            ("vendor/", "Bangun ulang dengan Composer", "Folder di-ignore dan harus cocok dengan composer.lock"),
            ("admin/tmp dan tmp", "Boleh bersihkan isi cache setelah backup", "Template terkompilasi dapat menyimpan versi lama"),
            ("plugins/logistik_non_medis", "Deploy dari tag rilis", "Kode utama modul yang sedang dikembangkan"),
            ("service-worker.js", "Periksa versi cache", "Cache browser dapat membuat tampilan seolah belum ter-update"),
        ],
        [3.8, 5.4, 7.5],
    )
    add_callout(doc, "Larangan database", "Jangan mengimpor mlite_db.sql atau mlite_only.sql ke database produksi yang sudah berisi data. File tersebut untuk instalasi awal, bukan update rutin.", "danger")

    add_title(doc, "4. Menyiapkan Commit Lokal yang Bersih")
    doc.add_paragraph("Lakukan di komputer pengembangan Windows/Laragon sebelum branch diganti atau dikirim ke server.")
    add_code(
        doc,
        """
cd C:\\laragon\\www\\mlite_rsns
git status --short
git diff --stat
git diff -- plugins/logistik_non_medis
""",
        "Langkah 1 - audit perubahan",
    )
    add_bullet(doc, "Pastikan file yang berubah memang bagian dari fitur Logistik Non Medis.")
    add_bullet(doc, "Jangan ikutkan agent/db_backups, dump SQL, log, credential, dan file hasil uji sementara.")
    add_bullet(doc, "Periksa config.php tetap tidak terlacak oleh Git.")
    add_code(
        doc,
        """
git add -p plugins/logistik_non_medis
git add service-worker.js
git diff --cached --stat
git diff --cached
git commit -m "feat(logistik_non_medis): finalisasi alur dan inventaris RSNS"
git status --short
""",
        "Langkah 2 - stage selektif dan commit",
    )
    add_callout(doc, "Mengapa git add -p", "Mode patch memungkinkan memilih perubahan sedikit demi sedikit. Ini jauh lebih aman pada Admin.php yang memuat banyak fitur sekaligus.", "info")
    add_checklist(
        doc,
        [
            "Tidak ada credential atau dump database di staged changes.",
            "File yang dibuat pengguna pada uploads tidak masuk commit.",
            "PHP dan JavaScript sudah lulus pemeriksaan sintaks.",
            "Pesan commit menjelaskan perubahan, bukan hanya 'update'.",
            "git status bersih sebelum pindah branch atau merge.",
        ],
    )

    add_title(doc, "5. Mengambil Update Baru dari Rifangga")
    doc.add_paragraph("Bab ini dipakai ketika git log HEAD..rifangga/main menampilkan commit baru.")
    add_code(
        doc,
        """
git fetch rifangga main --prune
git log --oneline --decorate HEAD..rifangga/main
git diff --name-status HEAD...rifangga/main
git switch main
git pull --ff-only origin main
git switch -c integrasi-rifangga-YYYYMMDD
git merge --no-ff --no-commit rifangga/main
""",
        "Buat branch integrasi khusus",
    )
    add_callout(doc, "Kondisi saat ini", "Pada 02 Agustus 2026, HEAD..rifangga/main kosong. Jika tetap kosong, jangan membuat merge kosong; lanjutkan ke pengujian dan penerbitan rilis lokal.", "success")

    doc.add_paragraph("Jika muncul konflik:")
    add_code(
        doc,
        """
git status
git diff --name-only --diff-filter=U
git diff --cc plugins/logistik_non_medis/Admin.php
# edit konflik secara manual
git add <file-yang-sudah-diselesaikan>
git diff --check
git commit
""",
        "Selesaikan konflik per file/fungsi",
    )
    add_table(
        doc,
        ["Area konflik", "Keputusan yang dipertahankan"],
        [
            ("Admin.php", "Alur status SPPB RSNS, transaksi stok, migrasi aman, Master Inventaris, dan hak akses lokal"),
            ("logistik.js", "Overlay proses, pencegah klik ganda, parsing JSON aman, dan endpoint terbaru"),
            ("logistik.css", "Tata letak responsif dan modal/form yang sudah diperbaiki"),
            ("view/admin/*.html", "Tab aktif/riwayat, label tambahan, breadcrumb, filter, PDF, QR, dan tampilan role"),
            ("service-worker.js", "Pengecualian halaman admin dan kenaikan nama cache bila aset statis berubah"),
            ("Skema database", "CREATE/ALTER idempotent; tidak menghapus tabel/kolom/data produksi"),
        ],
        [4.5, 12.2],
    )
    add_callout(doc, "Jangan memilih satu file utuh", "Hindari git checkout --ours atau --theirs untuk seluruh Admin.php. File tersebut berisi perubahan Rifangga dan penyesuaian RSNS yang sama-sama dibutuhkan.", "danger")

    add_title(doc, "6. Pemeriksaan Teknis Sebelum Rilis")
    add_code(
        doc,
        """
php -l plugins/logistik_non_medis/Admin.php
node --check plugins/logistik_non_medis/js/admin/logistik.js
git diff --check
composer validate --no-check-publish
composer install
""",
        "Pemeriksaan minimum",
    )
    doc.add_paragraph("Lakukan uji pada database salinan/staging, bukan langsung pada database produksi.")
    add_table(
        doc,
        ["Peran/Area", "Skenario smoke test", "Hasil yang diharapkan"],
        [
            ("Staf unit", "Buat SPPB rutin, tambahan mingguan, non rutin, dan retur", "Jenis dan status tampil konsisten; tidak membuat duplikat"),
            ("Kanit/Kasie/Kabid", "Buka antrean persetujuan dan setujui/tolak", "Hanya tahapan yang berhak yang dapat diproses"),
            ("Logistik", "Verifikasi, siapkan barang, konfirmasi pengambilan", "Stok berkurang sekali dan status menjadi selesai"),
            ("Gudang", "Lihat stok, mutasi 20 data per halaman, cetak kartu stok", "Urutan desc, filter dan PDF portrait benar"),
            ("Master", "Tambah/edit/hapus barang dan kategori", "JSON valid; dialog hapus dan overlay proses bekerja"),
            ("Aset", "Registrasi dari Kelompok > Jenis > Barang dan cetak QR", "KIB A-F lama tidak dipakai; nomor inventaris unik"),
            ("Navigasi", "Klik seluruh breadcrumb Logistik Non Medis", "Kembali ke menu yang sesuai, bukan selalu ke home"),
            ("Browser", "Ctrl+F5 dan uji desktop/mobile", "Tidak ada layout terpotong atau cache versi lama"),
        ],
        [3.0, 7.2, 6.5],
    )
    add_callout(doc, "Kriteria berhenti", "Jangan lanjut deploy bila ada warning PHP yang masuk ke respons JSON, SQLSTATE, status SPPB meloncat, stok berubah dua kali, atau fitur hapus/simpan belum memberi hasil final.", "danger")

    add_title(doc, "7. Menerbitkan Rilis ke Origin")
    doc.add_paragraph("Setelah branch kerja/integrasi lulus uji, gabungkan ke main milik RSNS.")
    add_code(
        doc,
        """
git switch main
git pull --ff-only origin main
git merge --no-ff fitur-hak-akses-style-lama
# atau merge branch integrasi-rifangga-YYYYMMDD
git push origin main
git tag -a logistik-v2026.08.02-1 -m "Rilis Logistik Non Medis RSNS"
git push origin logistik-v2026.08.02-1
git rev-parse --short HEAD
""",
        "Publish main dan tag",
    )
    add_callout(doc, "Tag adalah titik deploy", "Catat nama tag dan hash commit pada tiket/perubahan server. Server sebaiknya checkout tag yang sama agar versi yang dipasang tidak berubah diam-diam.", "success")

    add_title(doc, "8. Persiapan Server")
    add_checklist(
        doc,
        [
            "Jadwal maintenance sudah disetujui dan pengguna diberi pemberitahuan.",
            "PHP server sesuai composer.lock; proyek menetapkan platform PHP 7.4.33.",
            "Ekstensi PHP yang dibutuhkan mPDF, PDO MySQL, mbstring, GD, dan fileinfo tersedia.",
            "Composer tersedia atau vendor dibangun pada mesin dengan platform yang sama.",
            "Database user mempunyai hak CREATE/ALTER untuk migrasi idempotent modul.",
            "Ruang disk cukup untuk backup database, uploads, dan satu salinan kode.",
            "Port WebSocket 3892 dibuka hanya bila notifikasi realtime dipakai.",
        ],
    )
    add_code(
        doc,
        """
php -v
php -m | grep -E "pdo_mysql|mbstring|gd|fileinfo"
composer --version
git --version
df -h
""",
        "Pemeriksaan server Linux/VPS",
    )
    add_callout(doc, "Produksi", "Di config.php server gunakan DEV_MODE false dan jangan pernah menyalin config.php dari komputer Laragon. Nilai host, user, password, dan nama database produksi harus tetap milik server.", "warning")

    add_title(doc, "9. Backup Wajib Sebelum Deploy")
    add_code(
        doc,
        """
export RELEASE_TIME=$(date +%Y%m%d-%H%M%S)
mkdir -p ~/backup-mlite/$RELEASE_TIME
cd /var/www/mlite_rsns
git rev-parse HEAD > ~/backup-mlite/$RELEASE_TIME/commit-before.txt
mysqldump --single-transaction --routines --triggers \\
  -h DB_HOST -u DB_USER -p DB_NAME \\
  > ~/backup-mlite/$RELEASE_TIME/database.sql
tar -czf ~/backup-mlite/$RELEASE_TIME/files-persisten.tar.gz \\
  config.php uploads
""",
        "Contoh backup Linux",
    )
    add_bullet(doc, "Masukkan password database melalui prompt, bukan di command history.")
    add_bullet(doc, "Pastikan database.sql tidak kosong dan arsip uploads dapat dibuka.")
    add_bullet(doc, "Salin backup ke lokasi lain bila server hanya memiliki satu disk.")
    add_callout(doc, "Data RSNS tidak boleh dikosongkan", "Deploy kode tidak memerlukan penghapusan SPPB, mutasi, stok, aset, master inventaris, atau data operasional lain.", "danger")

    add_title(doc, "10. Deploy Rutin di Server (In-place)")
    doc.add_paragraph("Metode ini sesuai bila instalasi server saat ini sudah berupa Git working tree dan downtime singkat dapat diterima.")
    add_code(
        doc,
        """
cd /var/www/mlite_rsns
git status --porcelain
git fetch origin --tags --prune
git checkout main
git pull --ff-only origin main
git checkout logistik-v2026.08.02-1
composer install --no-dev --prefer-dist --optimize-autoloader
php -l plugins/logistik_non_medis/Admin.php
""",
        "Update kode ke tag rilis",
    )
    add_callout(doc, "Jika git status server tidak bersih", "Hentikan deploy. Simpan git diff dan cari asal perubahan server. Jangan menjalankan reset --hard sebelum perubahan tersebut dipahami dan dibackup.", "danger")
    doc.add_paragraph("Pastikan file persisten tetap tersedia:")
    add_code(
        doc,
        """
test -f config.php && echo "config OK"
test -d uploads && echo "uploads OK"
chmod -R u+rwX,g+rwX uploads admin/tmp tmp
""",
        "Izin direktori tulis",
    )
    add_callout(doc, "Shared hosting", "Jika server tidak memiliki Git/SSH, buat ZIP dari tag rilis, unggah hanya kode, lalu pertahankan config.php dan uploads server. Tetap lakukan backup database dan file sebelum ekstraksi.", "info")

    add_title(doc, "11. Migrasi Database Modul")
    doc.add_paragraph("Sebagian besar skema Logistik Non Medis dibuat atau diperbarui secara idempotent oleh fungsi _init... pada Admin.php ketika menu terkait dibuka.")
    add_number(doc, "Backup database produksi.")
    add_number(doc, "Pastikan user database aplikasi mempunyai izin CREATE TABLE dan ALTER TABLE.")
    add_number(doc, "Login sebagai administrator setelah kode terpasang.")
    add_number(doc, "Buka Dashboard Logistik, SPPB, Gudang Stok, Gudang Mutasi, Master Inventaris, dan Registrasi Aset satu kali.")
    add_number(doc, "Periksa log PHP/MySQL untuk warning ALTER, duplicate column, enum, atau permission denied.")
    add_number(doc, "Jalankan hakakses_migration.sql hanya bila tabel hak akses belum ada; skripnya menggunakan CREATE IF NOT EXISTS dan INSERT IGNORE.")
    add_callout(doc, "Rollback database tidak otomatis", "Mengembalikan commit Git tidak membatalkan perubahan skema. Karena itu backup sebelum migrasi wajib disimpan sampai rilis dinyatakan stabil.", "warning")

    add_title(doc, "12. Cache, Template, dan Service Worker")
    add_bullet(doc, "Bersihkan isi cache template admin/tmp dan tmp melalui fasilitas server/panel bila tampilan lama masih muncul.")
    add_bullet(doc, "Jangan menghapus foldernya; pastikan web server tetap dapat menulis ulang cache.")
    add_bullet(doc, "Naikkan cacheName pada service-worker.js saat aset statis yang di-cache berubah.")
    add_bullet(doc, "Pada browser uji gunakan Ctrl+F5 atau hapus site data/service worker bila masih memuat versi lama.")
    add_callout(doc, "Halaman admin", "service-worker.js saat ini melewati request /admin/ dan navigasi. Error FetchEvent pada halaman admin tetap harus diperiksa, tetapi tidak boleh menjadi alasan melakukan submit dua kali.", "info")

    add_title(doc, "13. Workerman dan Notifikasi Realtime")
    doc.add_paragraph("Notifikasi realtime menggunakan WebSocket pada port 3892. Fitur inti HTTP tetap harus diuji terpisah dari Workerman.")
    add_code(
        doc,
        """
cd /var/www/mlite_rsns
php workerman.php status
php workerman.php start -d
# setelah update kode worker
php workerman.php restart -d
""",
        "Perintah dasar Workerman",
    )
    add_bullet(doc, "Gunakan systemd/Supervisor pada server produksi agar worker otomatis hidup kembali.")
    add_bullet(doc, "Pastikan firewall/reverse proxy mengizinkan koneksi WebSocket 3892 bila port tersebut diekspos.")
    add_bullet(doc, "Jika koneksi gagal, aplikasi tidak boleh mengulang koneksi tanpa batas atau menggandakan proses simpan.")

    add_title(doc, "14. Smoke Test Setelah Deploy")
    add_checklist(
        doc,
        [
            "Halaman login dan dashboard tampil tanpa Notice/Warning PHP.",
            "Akun admin, staf unit, Kanit, Kasie, Kabid, dan logistik mendapat menu sesuai haknya.",
            "SPPB rutin dan tambahan memiliki label yang benar di sisi unit, admin, dan logistik.",
            "SPPB non rutin mengikuti persetujuan berjenjang dan tidak salah diarahkan ke alur rutin.",
            "Konfirmasi pengambilan hanya menyelesaikan status Siap Ambil/Siap Diserahkan.",
            "Stok dan kartu stok berubah tepat satu kali setelah transaksi.",
            "Gudang mutasi tampil desc, 20 item per halaman, tanpa kolom lokasi yang tidak dipakai.",
            "Filter kategori dan pagination gudang stok tidak menampilkan variabel undefined.",
            "PDF kartu stok portrait, tanpa identitas contoh mLITE Indonesia.",
            "Registrasi aset memakai Kelompok > Jenis > Barang terbaru dan QR dapat dibuka.",
            "Modal simpan/hapus menunggu respons final dan tombol terkunci selama proses.",
            "Console browser tidak memiliki error JSON Unexpected token '<' atau ReferenceError.",
        ],
    )
    add_callout(doc, "Observasi awal", "Pantau error log, log web server, dan transaksi selama minimal satu siklus kerja. Jangan hapus backup sebelum alur unit sampai logistik selesai diuji.", "success")

    add_title(doc, "15. Rollback Kode dan Database")
    doc.add_paragraph("Rollback dilakukan bila rilis menyebabkan gangguan yang tidak dapat diperbaiki cepat.")
    add_code(
        doc,
        """
cd /var/www/mlite_rsns
cat ~/backup-mlite/RELEASE_TIME/commit-before.txt
git fetch origin --tags
git checkout <COMMIT_SEBELUM_DEPLOY>
composer install --no-dev --prefer-dist --optimize-autoloader
php workerman.php restart -d
""",
        "Rollback kode",
    )
    add_callout(doc, "Jangan langsung restore database", "Jika migrasi hanya menambah tabel/kolom, versi kode lama biasanya masih dapat berjalan. Restore database akan menghilangkan transaksi yang masuk setelah backup. Putuskan bersama penanggung jawab data.", "warning")
    doc.add_paragraph("Jika restore database benar-benar diperlukan:")
    add_number(doc, "Hentikan akses pengguna dan Workerman.")
    add_number(doc, "Catat serta ekspor transaksi yang terjadi setelah backup.")
    add_number(doc, "Buat backup kondisi gagal sebelum restore.")
    add_number(doc, "Restore database.sql dengan persetujuan penanggung jawab.")
    add_number(doc, "Uji jumlah SPPB, mutasi, stok, aset, dan user setelah restore.")

    add_title(doc, "16. Troubleshooting Cepat")
    add_table(
        doc,
        ["Gejala", "Penyebab umum", "Tindakan"],
        [
            ("Unexpected token '<' pada JSON", "Notice/Warning PHP masuk sebelum JSON", "Buka Network response; perbaiki warning backend; jangan hanya menambah try/catch JS"),
            ("Halaman masih versi lama", "Template/browser/service worker cache", "Bersihkan cache template, Ctrl+F5, cek cacheName"),
            ("SQLSTATE / unknown column", "Migrasi belum berjalan atau DB user tanpa ALTER", "Buka menu pemicu migrasi dan periksa privilege/log"),
            ("WebSocket 3892 gagal", "Worker mati, firewall, atau proxy", "Cek status Workerman; uji port; pastikan fitur HTTP tetap berjalan"),
            ("git pull ditolak", "Server memiliki perubahan lokal", "Hentikan, simpan diff/backup, deploy dari tag pada working tree bersih"),
            ("Composer gagal", "Versi PHP/extension tidak cocok", "Samakan PHP dengan composer.lock dan pasang extension yang hilang"),
            ("Status SPPB salah", "Enum/mapping status lama dan baru tidak sinkron", "Periksa data aktual, migrasi status, dan fungsi approval sebelum mengubah massal"),
            ("Stok berubah dua kali", "Klik ganda/retry/endpoint dipanggil ulang", "Blokir proses, audit kartu stok dan log; jangan koreksi tanpa referensi transaksi"),
        ],
        [3.7, 5.4, 7.6],
    )

    add_title(doc, "17. Checklist Persetujuan Go-Live")
    add_table(
        doc,
        ["Tahap", "PIC", "Bukti", "Paraf"],
        [
            ("Review perubahan dan commit", "Developer/Rifangga", "Hash commit + diff stat", ""),
            ("Uji staging", "Developer + user perwakilan", "Checklist smoke test", ""),
            ("Backup database dan file", "Administrator server", "Lokasi + ukuran backup", ""),
            ("Deploy tag rilis", "Administrator server", "Tag + waktu deploy", ""),
            ("Migrasi dan cache", "Developer", "Log tanpa error", ""),
            ("Uji produksi", "Unit + Logistik", "Nomor transaksi uji", ""),
            ("Monitoring dan penutupan", "Penanggung jawab aplikasi", "Catatan observasi", ""),
        ],
        [4.7, 4.1, 5.4, 2.5],
    )

    add_title(doc, "18. Catatan Rilis")
    add_table(
        doc,
        ["Kolom", "Isi"],
        [
            ("Tag rilis", "........................................................"),
            ("Commit sebelum deploy", "........................................................"),
            ("Commit sesudah deploy", "........................................................"),
            ("Tanggal/jam", "........................................................"),
            ("Lokasi backup DB", "........................................................"),
            ("Lokasi backup uploads/config", "........................................................"),
            ("Pelaksana", "........................................................"),
            ("Hasil smoke test", "Lulus / Gagal / Catatan: ................................"),
            ("Keputusan", "Go-live / Rollback"),
        ],
        [5.2, 11.5],
    )

    add_callout(
        doc,
        "Kesimpulan",
        "Update Rifangga selalu masuk melalui branch integrasi. Rilis server selalu berasal dari origin RSNS dan tag yang sudah diuji. Dengan pemisahan ini, pembaruan upstream dapat diambil tanpa mengurangi fitur, konfigurasi, maupun data program yang sedang dipakai.",
        "success",
    )

    doc.core_properties.modified = datetime.now()
    doc.save(OUTPUT_FILE)
    return OUTPUT_FILE


if __name__ == "__main__":
    output = build_document()
    print(output)
