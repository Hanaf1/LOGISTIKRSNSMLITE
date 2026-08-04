from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUTPUT_DIR = Path(__file__).resolve().parent
DOC_DATE = "02 Agustus 2026"
NAVIGATION = "Menu > Logistik Non Medis"

NAVY = "34495E"
BLUE = "4A90E2"
GREEN = "70AD47"
ORANGE = "F39C34"
LIGHT_BLUE = "EAF3F8"
LIGHT_GREEN = "EDF6E8"
LIGHT_GRAY = "F2F3F4"
MID_GRAY = "D5D8DC"
WHITE = "FFFFFF"
TEXT = "404B55"
RED = "C0392B"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Halaman ")
    run.font.size = Pt(8)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def configure_document(doc, short_title):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.9)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color in (
        ("Title", 24, NAVY),
        ("Heading 1", 16, NAVY),
        ("Heading 2", 12.5, BLUE),
        ("Heading 3", 10.5, TEXT),
    ):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    header = section.header.paragraphs[0]
    header.text = f"PANDUAN LOGISTIK NON MEDIS  |  {short_title.upper()}"
    header.style = doc.styles["Normal"]
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.bold = True
    header.runs[0].font.color.rgb = RGBColor.from_string(NAVY)
    header.paragraph_format.space_after = Pt(2)
    p_pr = header._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:color"), BLUE)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    props = doc.core_properties
    props.subject = "Panduan penggunaan modul Logistik Non Medis"
    props.author = "Tim Sistem Informasi"
    props.keywords = "logistik non medis, SPPB, rutin, non rutin, retur"


def add_cover(doc, audience, subtitle):
    doc.add_paragraph("DOKUMEN INTERNAL", style="Subtitle").alignment = WD_ALIGN_PARAGRAPH.CENTER
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(36)

    title = doc.add_paragraph("MODUL PENGGUNAAN", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    product = doc.add_paragraph("LOGISTIK NON MEDIS")
    product.alignment = WD_ALIGN_PARAGRAPH.CENTER
    product.runs[0].font.name = "Arial"
    product.runs[0].font.size = Pt(20)
    product.runs[0].font.bold = True
    product.runs[0].font.color.rgb = RGBColor.from_string(BLUE)

    band = doc.add_table(rows=1, cols=1)
    band.alignment = WD_TABLE_ALIGNMENT.CENTER
    band.autofit = False
    band.columns[0].width = Cm(15.5)
    cell = band.cell(0, 0)
    set_cell_shading(cell, NAVY)
    set_cell_margins(cell, 300, 250, 300, 250)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(audience)
    run.font.name = "Arial"
    run.font.size = Pt(17)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(WHITE)

    sub = doc.add_paragraph(subtitle)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_before = Pt(12)
    sub.runs[0].font.size = Pt(11)

    doc.add_paragraph("\n\n")
    info = doc.add_table(rows=3, cols=2)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info.autofit = False
    info.columns[0].width = Cm(4)
    info.columns[1].width = Cm(9)
    values = (
        ("Versi dokumen", "1.0"),
        ("Tanggal", DOC_DATE),
        ("Akses utama", NAVIGATION),
    )
    for row, values_row in zip(info.rows, values):
        row.cells[0].text, row.cells[1].text = values_row
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        for cell in row.cells:
            set_cell_margins(cell)

    doc.add_page_break()


def add_intro(doc, purpose, scope, contents):
    doc.add_heading("Tentang Panduan", level=1)
    add_note(doc, "Tujuan", purpose, LIGHT_BLUE, BLUE)
    doc.add_heading("Ruang Lingkup", level=2)
    add_bullets(doc, scope)
    doc.add_heading("Isi Dokumen", level=2)
    for index, item in enumerate(contents, 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(item).bold = True
    add_note(
        doc,
        "Catatan tampilan",
        "Nama tombol dapat sedikit berbeda sesuai hak akses. Gunakan status pada tabel sebagai acuan utama. Saat lapisan 'Memproses data...' tampil, tunggu sampai proses dan pemuatan ulang selesai agar transaksi tidak terkirim dua kali.",
        LIGHT_GRAY,
        NAVY,
    )


def add_note(doc, title, text, fill=LIGHT_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 150, 180, 150, 180)
    p = cell.paragraphs[0]
    r = p.add_run(f"{title}: ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(accent)
    p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        if isinstance(item, tuple):
            head, body = item
            p.add_run(head).bold = True
            p.add_run(body)
        else:
            p.add_run(item)


def add_steps(doc, steps):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(1.2)
    table.columns[1].width = Cm(14.8)
    hdr = table.rows[0]
    hdr.cells[0].text = "No"
    hdr.cells[1].text = "Langkah"
    set_repeat_table_header(hdr)
    for cell in hdr.cells:
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string(WHITE)
    for index, step in enumerate(steps, 1):
        row = table.add_row()
        row.cells[0].text = str(index)
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if isinstance(step, tuple):
            head, body = step
            p = row.cells[1].paragraphs[0]
            p.add_run(head).bold = True
            p.add_run(body)
        else:
            row.cells[1].text = step
        if index % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, LIGHT_GRAY)
        for cell in row.cells:
            set_cell_margins(cell)
    doc.add_paragraph()


def add_matrix(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if widths:
        for column, width in zip(table.columns, widths):
            column.width = Cm(width)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for index, heading in enumerate(headers):
        hdr.cells[index].text = heading
        set_cell_shading(hdr.cells[index], NAVY)
        set_cell_margins(hdr.cells[index])
        for run in hdr.cells[index].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string(WHITE)
    for row_index, values in enumerate(rows, 1):
        row = table.add_row()
        for col_index, value in enumerate(values):
            row.cells[col_index].text = str(value)
            set_cell_margins(row.cells[col_index])
            if row_index % 2 == 0:
                set_cell_shading(row.cells[col_index], LIGHT_GRAY)
    doc.add_paragraph()


def add_section_break(doc):
    doc.add_section(WD_SECTION.NEW_PAGE)


def save(doc, filename, title):
    doc.core_properties.title = title
    target = OUTPUT_DIR / filename
    doc.save(target)
    return target


def build_staff_guide():
    doc = Document()
    configure_document(doc, "Staf Unit")
    add_cover(
        doc,
        "PANDUAN STAF UNIT",
        "Permintaan Rutin, Permintaan Non Rutin, dan Retur Barang",
    )
    add_intro(
        doc,
        "Membantu staf unit membuat, memantau, memperbaiki, dan menyelesaikan pengajuan barang tanpa mencampur alur Rutin, Non Rutin, serta Retur.",
        [
            "Pembuatan SPPB Rutin mingguan dan barang tambahan pada minggu yang sama.",
            "Pembuatan pengadaan Non Rutin tanpa kuota mingguan.",
            "Pembuatan Retur atas barang yang sebelumnya telah diserahkan.",
            "Pemantauan status melalui tab Aktif, Riwayat, dan detail pengajuan.",
        ],
        [
            "Mengenali menu dan status",
            "Permintaan Rutin pertama",
            "Barang tambahan Rutin",
            "Permintaan Non Rutin",
            "Retur barang",
            "Perbaikan masalah umum",
        ],
    )

    doc.add_heading("1. Masuk dan Membuka Menu", level=1)
    add_steps(doc, [
        "Masuk menggunakan akun staf unit yang telah diberikan hak akses Logistik Non Medis.",
        "Buka Menu, lalu pilih Logistik Non Medis.",
        "Pilih tab Distribusi untuk mengakses Permintaan BHP Mingguan, Pengadaan Non Rutin, dan Retur Barang.",
        "Pastikan nama pengguna dan unit yang tampil di bagian bawah aplikasi sudah benar.",
    ])
    add_note(doc, "Penting", "Unit Pengaju dan Diajukan Oleh diambil dari akun yang sedang masuk. Kolom yang terkunci tidak perlu diketik ulang.", LIGHT_GREEN, GREEN)

    doc.add_heading("2. Permintaan Rutin Mingguan", level=1)
    doc.add_paragraph("Gunakan Rutin untuk barang habis pakai yang memang dibutuhkan berulang oleh unit.")
    doc.add_heading("2.1 Membuat permintaan pertama minggu ini", level=2)
    add_steps(doc, [
        "Buka Distribusi > Permintaan BHP Mingguan.",
        "Pada tab Aktif, klik Tambah Permintaan apabila minggu berjalan belum memiliki SPPB Rutin.",
        "Periksa Tgl. SPPB, Unit Pengaju, dan Jenis Permintaan = Rutin.",
        "Pilih barang, isi Jumlah, Satuan, serta Catatan bila diperlukan, kemudian tambahkan ke daftar.",
        "Periksa ulang seluruh baris barang. Hindari barang ganda dan pastikan jumlah tidak nol.",
        "Klik Ajukan/Simpan. Tunggu indikator proses selesai dan daftar Aktif dimuat kembali.",
        "Buka Lihat Status untuk memantau proses sampai Siap Ambil dan Selesai.",
    ])
    add_note(doc, "Aturan minggu berjalan", "Satu unit mempunyai satu kelompok SPPB Rutin per minggu. Selama pengajuan masih aktif, gunakan Edit untuk memperbaiki item yang diperbolehkan. Setelah selesai, gunakan fitur barang tambahan, bukan membuat alur Non Rutin.", LIGHT_BLUE, BLUE)

    doc.add_heading("2.2 Menambahkan barang setelah SPPB selesai", level=2)
    add_steps(doc, [
        "Pada tab Aktif, cari panel 'Pengadaan rutin minggu ini sudah selesai'.",
        "Klik Atur Barang Tambahan atau Tambah Barang Minggu Ini.",
        "Halaman detail SPPB lama akan terbuka. Isi Barang, Jumlah, dan Catatan pada bagian Tambah Barang ke SPPB Ini.",
        "Klik Tambahkan. Item masuk ke SPPB minggu yang sama dan diberi tanda Tambahan/Rutin Tambahan.",
        "Gunakan Edit atau Hapus selama item tambahan masih berstatus Diajukan.",
        "Tunggu logistik menyetujui jumlah, menandai Siap Ambil, lalu lakukan pengambilan seperti biasa.",
    ])
    add_note(doc, "Cara mengenali tambahan", "Pada daftar unit, admin, dan logistik, item tambahan ditandai label 'Tambahan', kolom Tambahan berisi jumlah item tambahan, dan nomor SPPB tetap sama dengan pengajuan Rutin minggu tersebut.", LIGHT_GREEN, GREEN)

    doc.add_heading("2.3 Status Rutin dan tindakan staf", level=2)
    add_matrix(doc, ["Status", "Arti", "Tindakan staf unit"], [
        ("Diajukan", "Terkirim dan menunggu proses", "Periksa data; edit/hapus hanya jika tombol masih tersedia."),
        ("Proses Logistik", "Jumlah sedang diperiksa", "Tunggu hasil Qty ACC."),
        ("Siap Ambil", "Barang telah disiapkan", "Datang ke logistik dan sebutkan No. SPPB."),
        ("Selesai", "Pengambilan telah dicatat", "Periksa Riwayat; gunakan tambahan bila masih di minggu yang sama."),
        ("Ditolak/Dibatalkan", "Pengajuan tidak dilanjutkan", "Buka detail untuk membaca alasan dan koordinasikan perbaikan."),
    ], [3.0, 5.2, 7.8])

    add_section_break(doc)
    doc.add_heading("3. Permintaan Non Rutin", level=1)
    doc.add_paragraph("Gunakan Non Rutin untuk kebutuhan yang tidak termasuk distribusi BHP mingguan. Alur ini tidak memakai kuota mingguan dan memerlukan persetujuan bertahap.")
    add_steps(doc, [
        "Buka Distribusi > Pengadaan Non Rutin, lalu pilih Tambah Permintaan.",
        "Isi Latar Belakang dan Tujuan, Sasaran Kegunaan, serta Rencana Digunakan.",
        "Pilih Sifat Permintaan. Gunakan Urgent hanya bila kebutuhan memang mendesak dan dapat dijelaskan.",
        "Periksa Diajukan Oleh dan Ka Unit. Isi Penanggung Jawab 1; Penanggung Jawab 2 bersifat opsional.",
        "Tambahkan barang dari Master apabila tersedia. Untuk barang yang belum ada, gunakan Usulan Baru dan lengkapi spesifikasi.",
        "Isi jumlah, satuan, estimasi harga, dan catatan item. Periksa subtotal dan seluruh kelengkapan.",
        "Klik Ajukan. Pantau tab Aktif sampai seluruh tahapan selesai; pengajuan selesai berpindah ke Riwayat.",
    ])
    doc.add_heading("3.1 Urutan status Non Rutin", level=2)
    add_matrix(doc, ["Tahap", "Status/penanggung jawab", "Yang perlu dilakukan staf"], [
        ("1", "Diajukan", "Pastikan dokumen dan item lengkap."),
        ("2", "Disetujui Ka. Sie", "Menunggu persetujuan bidang berikutnya."),
        ("3", "Disetujui Kabid", "Menunggu pengelolaan administrasi/logistik."),
        ("4", "Kasie Umum / Verifikasi", "Siapkan klarifikasi bila diminta."),
        ("5", "Rekap dan konsultasi", "Pantau detail pengajuan."),
        ("6", "Pengajuan dana / Proses Pengadaan", "Tunggu proses pembelian/penyediaan."),
        ("7", "Siap Diserahkan", "Koordinasikan penerimaan barang."),
        ("8", "Selesai", "Periksa hasil pada Riwayat."),
    ], [1.3, 6.0, 8.7])
    add_note(doc, "Bedakan alur", "Rutin dan Rutin Tambahan tidak boleh diteruskan sebagai Non Rutin. Non Rutin hanya digunakan untuk pengadaan di luar distribusi rutin dan memiliki tahap persetujuan lebih panjang.", LIGHT_GRAY, ORANGE)

    doc.add_heading("4. Retur Barang", level=1)
    doc.add_paragraph("Retur dipakai untuk mengembalikan barang dari unit berdasarkan SPPB yang telah diserahkan.")
    add_steps(doc, [
        "Buka Distribusi > Retur Barang, lalu klik Tambah Retur.",
        "Pilih nomor SPPB sumber. Sistem memuat barang dan jumlah yang pernah diserahkan.",
        "Isi Qty Retur. Nilai tidak boleh melebihi Qty Kirim.",
        "Pilih alasan: Sisa, Salah Kirim, atau Rusak. Tambahkan Kondisi Fisik secara jelas.",
        "Simpan Draft. Selama status Pending, data masih dapat diedit atau dihapus bila tombol tersedia.",
        "Serahkan barang fisik kepada petugas logistik untuk inspeksi.",
        "Pantau status: Pending, Disetujui, atau Ditolak. Stok bertambah kembali hanya bila logistik menyetujui retur yang layak masuk stok.",
    ])
    add_note(doc, "Barang rusak", "Barang rusak tidak otomatis menjadi stok siap pakai. Tuliskan kondisi dengan jujur agar logistik dapat menentukan penanganan sesuai hasil inspeksi.", LIGHT_GRAY, RED)

    doc.add_heading("5. Pencarian, Aktif, dan Riwayat", level=1)
    add_bullets(doc, [
        ("Tab Aktif: ", "menampilkan pengajuan yang belum selesai dan masih membutuhkan tindakan."),
        ("Tab Riwayat: ", "menampilkan pengajuan yang sudah selesai, ditolak, atau dibatalkan."),
        ("Pencarian: ", "gunakan No. SPPB, nama unit, atau nama barang sesuai kolom yang tersedia."),
        ("Lihat Status/Detail: ", "gunakan untuk melihat Qty Minta, Qty ACC, status item, catatan, dan jejak tahapan."),
    ])

    doc.add_heading("6. Masalah Umum", level=1)
    add_matrix(doc, ["Masalah", "Pemeriksaan", "Tindakan"], [
        ("Tombol simpan seperti tidak bereaksi", "Ada lapisan Memproses data", "Tunggu sampai daftar dimuat ulang; jangan klik dua kali."),
        ("Tidak dapat edit", "Status sudah diproses atau selesai", "Gunakan detail untuk melihat status; hubungi logistik bila koreksi diperlukan."),
        ("Tambah malah membuat pengajuan baru", "Masuk dari tombol Tambah Permintaan", "Untuk tambahan Rutin, masuk lewat Atur Barang Tambahan pada SPPB selesai."),
        ("Barang tidak tersedia", "Barang belum ada pada master", "Pada Non Rutin pilih Usulan Baru; pada Rutin hubungi pengelola master barang."),
        ("Status belum berubah", "Halaman belum termuat ulang", "Tunggu proses, lalu gunakan Refresh satu kali."),
    ], [4.2, 5.2, 6.6])

    doc.add_heading("Checklist Staf Unit", level=1)
    add_bullets(doc, [
        "Akun dan unit pengaju sudah benar.",
        "Jenis pengajuan dipilih sesuai kebutuhan: Rutin, Rutin Tambahan, Non Rutin, atau Retur.",
        "Barang, jumlah, satuan, dan catatan sudah diperiksa.",
        "Nomor SPPB/Retur disimpan sebagai referensi komunikasi.",
        "Status dipantau sampai Siap Ambil/Siap Diserahkan dan Selesai.",
    ])
    return save(doc, "01_Modul_Staf_Unit_Logistik_Non_Medis.docx", "Modul Staf Unit - Logistik Non Medis")


def build_approval_guide():
    doc = Document()
    configure_document(doc, "Kanit, KASI, KABID")
    add_cover(
        doc,
        "PANDUAN PENANGGUNG JAWAB DAN PEMBERI PERSETUJUAN",
        "Kepala Unit (Kanit), Kepala Seksi (KASI), dan Kepala Bidang (KABID)",
    )
    add_intro(
        doc,
        "Menyeragamkan pemeriksaan kebutuhan, pemberian persetujuan, penolakan, serta pemantauan pengajuan Non Rutin sesuai kewenangan masing-masing.",
        [
            "Pemeriksaan identitas unit, tujuan, penanggung jawab, barang, jumlah, dan estimasi biaya.",
            "Persetujuan bertahap oleh pejabat sesuai hak akses dan status pengajuan.",
            "Penolakan dengan alasan yang dapat ditindaklanjuti.",
            "Pemantauan Rutin, Rutin Tambahan, Non Rutin, dan Retur tanpa mencampur kewenangan.",
        ],
        [
            "Prinsip dan pembagian peran",
            "Pemeriksaan sebelum menyetujui",
            "Langkah persetujuan KASI dan KABID",
            "Peran Kanit",
            "Status dan kontrol internal",
        ],
    )

    doc.add_heading("1. Prinsip Persetujuan", level=1)
    add_bullets(doc, [
        "Setujui berdasarkan kebutuhan yang dapat dijelaskan, bukan hanya karena form sudah terisi.",
        "Pastikan jumlah, spesifikasi, waktu penggunaan, dan penanggung jawab konsisten.",
        "Gunakan akun sendiri. Jangan meminjamkan akun karena tindakan tersimpan sebagai jejak audit.",
        "Jangan menyetujui kembali Rutin sebagai Non Rutin. Kedua jenis memiliki tujuan dan alur berbeda.",
        "Beri alasan yang spesifik saat menolak agar staf unit dapat memperbaiki pengajuan.",
    ])

    doc.add_heading("2. Pembagian Peran", level=1)
    add_matrix(doc, ["Peran", "Fokus pemeriksaan", "Tindakan utama"], [
        ("Kanit / Kepala Unit", "Kebutuhan unit, kelengkapan, penanggung jawab, dan kesesuaian rencana penggunaan", "Memeriksa dan memantau pengajuan unit; melakukan persetujuan bila tombol tersedia sesuai konfigurasi hak akses."),
        ("KASI", "Kelayakan operasional, urgensi, jumlah, dan spesifikasi", "Memberikan persetujuan formal pertama pada Non Rutin berstatus Diajukan."),
        ("KABID", "Prioritas bidang, kewajaran biaya, risiko, dan persetujuan KASI", "Memberikan persetujuan lanjutan setelah KASI."),
        ("Logistik", "Ketersediaan, verifikasi teknis, pengadaan, distribusi, dan stok", "Melanjutkan proses setelah persetujuan struktural lengkap."),
    ], [3.0, 6.1, 7.0])
    add_note(doc, "Hak akses", "Tombol persetujuan hanya tampil ketika akun memiliki peran yang tepat dan status pengajuan berada pada tahapnya. Apabila tombol tidak tampil, jangan memakai akun peran lain; minta administrator memeriksa pemetaan akun.", LIGHT_BLUE, BLUE)

    doc.add_heading("3. Pemeriksaan Sebelum Persetujuan", level=1)
    add_steps(doc, [
        "Masuk ke Logistik Non Medis > Distribusi > Pengadaan Non Rutin.",
        "Buka tab Aktif dan cari pengajuan berdasarkan No. SPPB atau unit.",
        "Buka Detail/Lihat Status dan cocokkan Unit Pengaju, Diajukan Oleh, Ka Unit, serta Penanggung Jawab.",
        "Baca Latar Belakang dan Tujuan, Sasaran Kegunaan, Rencana Digunakan, dan Sifat Permintaan.",
        "Periksa setiap item: nama/spesifikasi, jumlah, satuan, estimasi harga, subtotal, dan catatan.",
        "Pastikan status saat ini memang menunggu tindakan peran Anda.",
        "Pilih Setujui untuk meneruskan atau Tolak dengan alasan yang jelas. Tunggu proses selesai.",
        "Buka ulang detail untuk memastikan status dan nama pemberi persetujuan telah tercatat.",
    ])

    doc.add_heading("4. Alur Persetujuan Non Rutin", level=1)
    add_matrix(doc, ["Urutan", "Status sebelum tindakan", "Pelaksana", "Hasil"], [
        ("1", "Diajukan", "Kanit memeriksa kelengkapan unit", "Siap dinilai KASI / disetujui bila hak akses dikonfigurasi"),
        ("2", "Diajukan", "KASI", "Disetujui Ka. Sie atau Ditolak"),
        ("3", "Disetujui Ka. Sie", "KABID", "Disetujui Kabid atau Ditolak"),
        ("4", "Disetujui Kabid", "Kasie Umum / Logistik", "Masuk verifikasi dan proses pengadaan"),
    ], [1.5, 4.6, 3.5, 6.5])

    doc.add_heading("4.1 Langkah KASI", level=2)
    add_steps(doc, [
        "Filter daftar pada status Diajukan.",
        "Buka detail dan pastikan Kanit/penanggung jawab unit telah melakukan pemeriksaan internal.",
        "Nilai apakah kebutuhan sesuai fungsi unit dan apakah sifat Urgent dapat dipertanggungjawabkan.",
        "Periksa usulan barang baru lebih teliti karena belum memiliki referensi master.",
        "Klik Setujui untuk menghasilkan status Disetujui Ka. Sie, atau Tolak dengan alasan koreksi.",
    ])

    doc.add_heading("4.2 Langkah KABID", level=2)
    add_steps(doc, [
        "Filter daftar pada status Disetujui Ka. Sie.",
        "Periksa persetujuan KASI, tujuan, prioritas, total estimasi, dan dampak terhadap kebutuhan bidang.",
        "Pastikan pengajuan tidak memecah kebutuhan yang sama menjadi beberapa SPPB tanpa alasan.",
        "Klik Setujui untuk menghasilkan status Disetujui Kabid, atau Tolak dengan alasan yang dapat ditindaklanjuti.",
    ])

    doc.add_heading("5. Posisi Rutin, Tambahan, dan Retur", level=1)
    add_matrix(doc, ["Jenis", "Cara mengenali", "Peran Kanit/KASI/KABID"], [
        ("Rutin", "Label Rutin; kebutuhan mingguan", "Kanit mengawasi kewajaran kebutuhan unit. Proses operasional dilanjutkan logistik."),
        ("Rutin Tambahan", "Label Tambahan/Rutin Tambahan dan nomor SPPB mingguan yang sama", "Pastikan tambahan benar-benar kebutuhan minggu berjalan; bukan pengadaan Non Rutin terselubung."),
        ("Non Rutin", "Form tujuan, sasaran, rencana, penanggung jawab, dan estimasi biaya", "KASI dan KABID memberikan persetujuan formal bertahap."),
        ("Retur", "Nomor Retur terhubung ke SPPB yang telah diserahkan", "Kanit mengawasi pertanggungjawaban unit; inspeksi dan keputusan stok dilakukan logistik."),
    ], [3.0, 6.0, 7.1])

    doc.add_heading("6. Kriteria Setujui atau Tolak", level=1)
    add_matrix(doc, ["Setujui bila", "Tolak/kembalikan bila"], [
        ("Tujuan jelas dan sesuai fungsi unit", "Tujuan kosong, umum, atau tidak relevan"),
        ("Jumlah dan spesifikasi wajar", "Jumlah tidak dapat dijelaskan atau spesifikasi tidak cukup"),
        ("Penanggung jawab tersedia", "Penanggung jawab tidak jelas"),
        ("Estimasi biaya dapat ditelusuri", "Harga/subtotal tidak wajar atau tidak lengkap"),
        ("Tidak menduplikasi pengajuan aktif", "Ada pengajuan aktif untuk kebutuhan yang sama"),
    ], [8.0, 8.0])

    doc.add_heading("7. Status dan Arti", level=1)
    add_matrix(doc, ["Status", "Arti"], [
        ("Diajukan", "Menunggu persetujuan awal."),
        ("Disetujui Ka. Sie", "KASI telah menyetujui; menunggu KABID."),
        ("Disetujui Kabid", "Persetujuan struktural selesai; masuk proses umum/logistik."),
        ("Verifikasi/Proses Pengadaan", "Sedang diperiksa atau dipenuhi oleh pengelola."),
        ("Siap Diserahkan", "Barang siap diserahkan kepada unit."),
        ("Selesai", "Serah terima sudah dicatat."),
        ("Ditolak/Dibatalkan", "Tidak dilanjutkan; alasan harus diperiksa."),
    ], [5.0, 11.0])

    doc.add_heading("Checklist Pemberi Persetujuan", level=1)
    add_bullets(doc, [
        "Akun dan peran yang digunakan sudah benar.",
        "Status memang menunggu tindakan peran saya.",
        "Tujuan, sasaran, rencana, dan penanggung jawab telah diperiksa.",
        "Barang, jumlah, spesifikasi, dan estimasi biaya masuk akal.",
        "Keputusan serta alasan penolakan dapat dipertanggungjawabkan.",
        "Status baru telah diperiksa setelah proses selesai.",
    ])
    return save(doc, "02_Modul_Persetujuan_Kanit_KASI_KABID.docx", "Modul Persetujuan Kanit KASI KABID - Logistik Non Medis")


def build_logistics_guide():
    doc = Document()
    configure_document(doc, "Petugas Logistik")
    add_cover(
        doc,
        "PANDUAN PETUGAS LOGISTIK",
        "Pengelolaan SPPB, Distribusi, Retur, Stok, Mutasi, dan Kartu Stok",
    )
    add_intro(
        doc,
        "Membantu petugas logistik memproses permintaan secara konsisten, menjaga ketepatan stok, serta mempertahankan jejak transaksi dari pengajuan sampai serah terima.",
        [
            "Proses Rutin dan Rutin Tambahan sampai pengambilan.",
            "Kelanjutan pengadaan Non Rutin setelah persetujuan struktural.",
            "Inspeksi dan keputusan Retur.",
            "Pemantauan stok, mutasi, kartu stok, master barang, dan audit transaksi.",
        ],
        [
            "Persiapan dan kontrol harian",
            "Memproses Rutin dan Tambahan",
            "Memproses Non Rutin",
            "Memproses Retur",
            "Stok, Mutasi, dan Kartu Stok",
            "Kontrol kesalahan dan rekonsiliasi",
        ],
    )

    doc.add_heading("1. Persiapan Kerja", level=1)
    add_steps(doc, [
        "Masuk dengan akun logistik/administrator yang memiliki hak proses transaksi.",
        "Buka dashboard Logistik Non Medis dan periksa kartu ringkasan permintaan, stok kritis, dan pekerjaan aktif.",
        "Gunakan kategori menu yang sesuai: Master Data, Pengadaan, Manajemen Gudang, Distribusi, Aset, atau Laporan & Audit.",
        "Periksa notifikasi serta tab Aktif sebelum membuka Riwayat.",
        "Pastikan master barang, satuan, kategori, dan saldo awal telah benar sebelum distribusi.",
    ])
    add_note(doc, "Kontrol klik ganda", "Setiap simpan, hapus, setujui, atau proses stok menampilkan indikator 'Memproses data...'. Jangan menutup halaman atau mengulangi klik sampai data selesai dimuat kembali.", LIGHT_BLUE, BLUE)

    doc.add_heading("2. Rutin dan Rutin Tambahan", level=1)
    doc.add_heading("2.1 Mengenali jenis transaksi", level=2)
    add_matrix(doc, ["Jenis", "Tanda di daftar", "Perlakuan"], [
        ("Rutin", "Label Rutin dan nomor SPPB mingguan", "Proses seluruh item aktif pada pengajuan awal."),
        ("Rutin Tambahan", "Label Tambahan/Rutin Tambahan; kolom Tambahan berisi item", "Proses hanya item tambahan yang berstatus aktif; item Rutin lama yang sudah selesai tidak diproses ulang."),
    ], [3.0, 5.5, 7.6])

    doc.add_heading("2.2 Memeriksa dan menyetujui jumlah", level=2)
    add_steps(doc, [
        "Buka Distribusi > Permintaan BHP Mingguan > tab Aktif.",
        "Gunakan filter status atau pencarian No. SPPB/unit/barang.",
        "Klik aksi Proses Logistik pada baris yang berstatus Diajukan.",
        "Periksa setiap item dan masukkan Qty ACC. Nilai tidak boleh melebihi kebutuhan tanpa alasan yang dapat dipertanggungjawabkan.",
        "Untuk Rutin Tambahan, pastikan tabel hanya memuat item berlabel Tambahan yang belum selesai.",
        "Klik Setujui & Tandai Siap Ambil. Tunggu daftar dimuat ulang.",
        "Buka detail dan pastikan status berubah menjadi Siap Ambil serta Qty ACC tersimpan.",
    ])

    doc.add_heading("2.3 Serah terima dan penyelesaian", level=2)
    add_steps(doc, [
        "Saat unit datang, cocokkan No. SPPB, unit, nama barang, dan Qty ACC.",
        "Klik aksi Pengambilan/Konfirmasi pada baris berstatus Siap Ambil atau Siap Diserahkan.",
        "Isi nama pengambil barang dengan identitas yang dapat ditelusuri.",
        "Simpan Pengambilan. Sistem mengubah status menjadi Selesai dan membukukan barang keluar.",
        "Periksa Mutasi/Kartu Stok menggunakan No. SPPB sebagai referensi.",
    ])
    add_note(doc, "Khusus tambahan", "Pembukuan item tambahan memakai referensi tambahan tersendiri di belakang No. SPPB. Ini mencegah stok item Rutin lama dikurangi untuk kedua kalinya.", LIGHT_GREEN, GREEN)

    doc.add_heading("3. Pengadaan Non Rutin", level=1)
    doc.add_paragraph("Non Rutin tidak memakai kuota mingguan. Petugas logistik melanjutkan proses setelah persetujuan KASI dan KABID sesuai status.")
    add_steps(doc, [
        "Buka Distribusi > Pengadaan Non Rutin dan pilih tab Aktif.",
        "Pastikan status telah melewati Disetujui Ka. Sie dan Disetujui Kabid sebelum melanjutkan tahap logistik.",
        "Periksa identitas pengajuan, tujuan, penanggung jawab, spesifikasi, jumlah, dan estimasi harga.",
        "Lakukan tahap Kasie Umum/Verifikasi sesuai kewenangan akun.",
        "Susun rekap Logistik Umum dan catat hasil konsultasi Kabid Umum bila diperlukan.",
        "Lanjutkan Pengajuan Dana/Bendahara dan Proses Pengadaan sesuai bukti administrasi.",
        "Setelah barang siap, ubah status menjadi Siap Diserahkan.",
        "Lakukan serah terima, catat penerima, lalu selesaikan pengajuan.",
    ])
    doc.add_heading("3.1 Urutan kendali", level=2)
    add_matrix(doc, ["Tahap", "Kontrol minimum"], [
        ("Verifikasi", "Persetujuan lengkap, item jelas, tidak duplikat, dan unit dapat dihubungi."),
        ("Rekap", "Jumlah dan estimasi konsisten dengan detail item."),
        ("Dana", "Nominal dan dokumen pendukung sesuai."),
        ("Pengadaan", "Sumber/vendor dan hasil pemenuhan dapat ditelusuri."),
        ("Penyerahan", "Barang, jumlah, penerima, waktu, dan status tercatat."),
    ], [4.0, 12.0])

    add_section_break(doc)
    doc.add_heading("4. Retur Barang dari Unit", level=1)
    add_steps(doc, [
        "Buka Distribusi > Retur Barang dan cari nomor Retur/SPPB/unit.",
        "Buka Inspeksi & Approval untuk retur berstatus Pending.",
        "Cocokkan barang fisik dengan No. SPPB, batch, Qty Kirim, Qty Retur, alasan, dan kondisi fisik.",
        "Isi Catatan Inspeksi Logistik untuk setiap item.",
        "Pilih Setujui Retur & Update Stok bila barang layak diterima sesuai kebijakan, atau Tolak Retur bila tidak sesuai.",
        "Tunggu proses selesai, lalu pastikan status menjadi Disetujui atau Ditolak.",
        "Bila disetujui, periksa mutasi masuk dan saldo akhir. Bila ditolak, pastikan alasan disampaikan kepada unit.",
    ])
    add_matrix(doc, ["Status Retur", "Arti", "Dampak stok"], [
        ("Pending", "Menunggu inspeksi/keputusan", "Belum berubah"),
        ("Disetujui", "Retur diterima", "Diperbarui sesuai hasil proses"),
        ("Ditolak", "Retur tidak diterima", "Tidak bertambah"),
    ], [4.0, 6.0, 6.0])
    add_note(doc, "Hindari koreksi manual", "Jangan membuat penyesuaian manual untuk menggantikan retur yang belum disetujui. Selesaikan alur Retur agar referensi transaksi dan catatan inspeksi tetap utuh.", LIGHT_GRAY, RED)

    doc.add_heading("5. Pengelolaan Stok", level=1)
    doc.add_heading("5.1 Daftar stok", level=2)
    add_bullets(doc, [
        "Gunakan pencarian nama/kode barang dan Filter Kategori untuk mempersempit daftar.",
        "Periksa Stok Akhir, Satuan, Harga Rata-rata, Total Nilai, dan Status stok.",
        "Gunakan Tambah Data hanya untuk master/saldo yang memang belum tersedia.",
        "Gunakan kartu stok pada item untuk menelusuri sumber pergerakan.",
        "Lakukan Export Excel bila dibutuhkan untuk rekonsiliasi dan dokumentasi.",
    ])

    doc.add_heading("5.2 Mutasi gudang", level=2)
    add_steps(doc, [
        "Buka Manajemen Gudang > Mutasi.",
        "Daftar awal menampilkan seluruh mutasi terbaru secara menurun, 20 data per halaman.",
        "Gunakan Filter Barang dan rentang tanggal bila ingin menelusuri transaksi tertentu.",
        "Cocokkan Tanggal, Barang, No. Referensi, Tipe, Batch, Masuk, Keluar, Stok Akhir, dan User.",
        "Gunakan Tambah Penyesuaian Manual hanya setelah stok fisik diverifikasi dan alasan dicatat.",
    ])

    doc.add_heading("5.3 Cetak kartu stok PDF", level=2)
    add_steps(doc, [
        "Pada Mutasi, pilih barang yang ingin dicetak.",
        "Isi Dari Tanggal dan Sampai Tanggal bila membutuhkan periode tertentu; kosongkan untuk semua tanggal.",
        "Klik Tampilkan dan periksa urutan mutasi terbaru ke lama.",
        "Pilih Cetak Kartu Stok PDF dari opsi cetak/ekspor yang tersedia.",
        "Pastikan dokumen portrait memuat barang, periode, transaksi, total masuk, total keluar, dan stok akhir.",
    ])

    doc.add_heading("6. Master dan Integritas Data", level=1)
    add_matrix(doc, ["Objek", "Yang harus dijaga", "Larangan"], [
        ("Master barang", "Kode unik, nama jelas, kategori dan satuan benar", "Jangan membuat barang duplikat dengan ejaan berbeda."),
        ("Stok", "Saldo sesuai transaksi dan hasil fisik", "Jangan mengubah saldo tanpa referensi/alasan."),
        ("SPPB", "Jenis, item, Qty ACC, status, dan penerima", "Jangan menandai selesai sebelum serah terima."),
        ("Retur", "SPPB sumber, qty, alasan, kondisi, dan inspeksi", "Jangan menambah stok sebelum disetujui."),
        ("Aset", "Identitas, unit, kelompok, kondisi, dan inventaris", "Jangan menghapus data aset untuk membersihkan transaksi BHP."),
    ], [3.0, 6.2, 6.8])

    doc.add_heading("7. Pemeriksaan Bila Terjadi Error", level=1)
    add_steps(doc, [
        "Catat halaman, No. SPPB/Retur, status, waktu, akun, dan tindakan terakhir.",
        "Tunggu indikator proses selesai, lalu Refresh satu kali.",
        "Periksa apakah transaksi sebenarnya sudah tersimpan agar tidak membuat duplikasi.",
        "Buka detail dan Mutasi untuk memeriksa konsistensi status dan stok.",
        "Jika masih bermasalah, simpan tangkapan layar pesan error dan laporkan tanpa mengubah data secara manual.",
    ])
    add_note(doc, "Pesan jaringan", "Kegagalan WebSocket lokal atau service worker tidak selalu berarti transaksi gagal. Konfirmasi hasil pada daftar dan database melalui tampilan aplikasi sebelum mengulang proses.", LIGHT_GRAY, ORANGE)

    doc.add_heading("Checklist Penutupan Harian", level=1)
    add_bullets(doc, [
        "Tidak ada proses yang berhenti pada Diajukan karena belum diperiksa.",
        "Semua serah terima hari ini memiliki nama pengambil/penerima.",
        "Mutasi keluar/masuk sesuai dengan SPPB dan Retur yang diselesaikan.",
        "Tidak ada stok negatif yang belum memiliki penjelasan dan tindak lanjut.",
        "Penyesuaian manual memiliki alasan serta petugas yang dapat ditelusuri.",
        "Pengajuan selesai masuk Riwayat dan pengajuan aktif tetap berada di tab Aktif.",
    ])
    return save(doc, "03_Modul_Petugas_Logistik_Non_Medis.docx", "Modul Petugas Logistik - Logistik Non Medis")


def validate_document(path):
    loaded = Document(path)
    if len(loaded.paragraphs) < 20 or len(loaded.tables) < 3:
        raise RuntimeError(f"Dokumen tidak lengkap: {path.name}")
    if not any("LOGISTIK NON MEDIS" in paragraph.text for paragraph in loaded.paragraphs):
        raise RuntimeError(f"Judul tidak ditemukan: {path.name}")
    return {
        "file": path.name,
        "paragraphs": len(loaded.paragraphs),
        "tables": len(loaded.tables),
        "bytes": path.stat().st_size,
    }


if __name__ == "__main__":
    generated = [build_staff_guide(), build_approval_guide(), build_logistics_guide()]
    for report in map(validate_document, generated):
        print(f"OK {report['file']} | {report['paragraphs']} paragraf | {report['tables']} tabel | {report['bytes']} byte")
